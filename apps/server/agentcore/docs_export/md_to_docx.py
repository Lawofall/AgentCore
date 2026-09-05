"""Markdown → DOCX deterministic converter (markdown-it-py + python-docx).

MVP coverage: headings #–####, paragraphs, ordered/unordered lists, tables,
fenced code, relative-path images (embedded), links. Missing images produce
explicit warnings (never silent). Default styles target a clean Chinese
公文-like look (黑体 headings / 宋体 body with Latin fallbacks).

段落几何按 ``layout`` 档位走（见 ``docs_export.layout``）：两档都是 A4、正文/列表
1.5 倍行距、一级标题居中；首行缩进两字、两端对齐、公文页边距、页码只在
``official`` 档开——档位来自调用方入参，绝不看正文内容猜。
"""

from __future__ import annotations

import io
import posixpath
import re
from collections.abc import Mapping
from dataclasses import dataclass, field
from typing import TYPE_CHECKING, Any
from urllib.parse import unquote, urlparse

from markdown_it import MarkdownIt
from markdown_it.token import Token

from agentcore.docs_export.layout import (
    A4_HEIGHT_CM,
    A4_WIDTH_CM,
    BODY_SPACE_AFTER_PT,
    FIRST_LINE_INDENT_CHARS,
    LAYOUT_OFFICIAL,
    LAYOUT_STANDARD,
    LINE_SPACING_MULTIPLE,
    OFFICIAL_MARGINS_CM,
    STANDARD_MARGINS_CM,
    DocLayout,
)

if TYPE_CHECKING:
    from docx.document import Document

# CommonMark + GFM tables. html=False keeps raw HTML out of the tree.
_MD = MarkdownIt("commonmark", {"html": False, "linkify": False, "breaks": False}).enable(
    "table"
)

_SCHEME_RE = re.compile(r"^[a-zA-Z][a-zA-Z0-9+.-]*:")

# 干净公文 defaults — Word substitutes when the OS lacks these faces.
_FONT_LATIN = "Times New Roman"
_FONT_BODY_CJK = "宋体"
_FONT_HEADING_CJK = "黑体"
_FONT_CODE = "Consolas"

_HEADING_PT = {1: 22, 2: 18, 3: 16, 4: 14}
_BODY_PT = 12
_CODE_PT = 10
_MAX_IMAGE_WIDTH_IN = 5.8

# python-docx is loaded on first convert (not at import). Sidecar / cloud both ship
# the Office stack; lazy load keeps tool registration from hard-crashing chat if a
# runtime is mis-bundled — capability still belongs in the dependency matrix.
_DocumentFactory: Any = None
_WD_ALIGN_PARAGRAPH: Any = None
_RT: Any = None
_qn: Any = None
_OxmlElement: Any = None
_Cm: Any = None
_Inches: Any = None
_Pt: Any = None
_RGBColor: Any = None
_MAX_IMAGE_WIDTH: Any = None


def _ensure_docx() -> None:
    """Bind python-docx symbols used by the converter (idempotent)."""
    global _DocumentFactory, _WD_ALIGN_PARAGRAPH, _RT, _qn, _OxmlElement
    global _Cm, _Inches, _Pt, _RGBColor, _MAX_IMAGE_WIDTH
    if _DocumentFactory is not None:
        return
    from docx import Document as DocumentFactory
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement
    from docx.shared import Cm, Inches, Pt, RGBColor

    _DocumentFactory = DocumentFactory
    _WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH
    _RT = RT
    _qn = qn
    _OxmlElement = OxmlElement
    _Cm = Cm
    _Inches = Inches
    _Pt = Pt
    _RGBColor = RGBColor
    _MAX_IMAGE_WIDTH = Inches(_MAX_IMAGE_WIDTH_IN)


@dataclass(frozen=True)
class MdToDocxResult:
    """Bytes of a .docx plus non-fatal export warnings (e.g. missing images)."""

    docx_bytes: bytes
    warnings: list[str] = field(default_factory=list)


def docx_path_for_markdown(md_path: str) -> str:
    """``报告.md`` → ``报告.docx`` (same directory)."""
    md_path = md_path.replace("\\", "/").strip()
    lower = md_path.lower()
    if lower.endswith(".markdown"):
        return md_path[: -len(".markdown")] + ".docx"
    if lower.endswith(".md"):
        return md_path[: -len(".md")] + ".docx"
    return f"{md_path}.docx"


def collect_image_srcs(markdown: str) -> list[str]:
    """Return unique image ``src`` values in document order."""
    seen: set[str] = set()
    out: list[str] = []
    for token in _MD.parse(markdown or ""):
        if token.type != "inline" or not token.children:
            continue
        for child in token.children:
            if child.type != "image":
                continue
            src = str(child.attrGet("src") or "").strip()
            if not src or src in seen:
                continue
            seen.add(src)
            out.append(src)
    return out


def is_embeddable_relative_src(src: str) -> bool:
    """True when ``src`` looks like a workspace-relative path (not a URL/data URI)."""
    s = (src or "").strip()
    if not s or s.startswith("#"):
        return False
    if s.startswith("data:"):
        return False
    return not bool(_SCHEME_RE.match(s))


def resolve_workspace_image_path(md_path: str, src: str) -> str | None:
    """Map a relative image src to a workspace-relative path, or None if unsafe/remote."""
    if not is_embeddable_relative_src(src):
        return None
    cleaned = unquote(src.strip()).replace("\\", "/")
    while cleaned.startswith("./"):
        cleaned = cleaned[2:]
    md_dir = posixpath.dirname(md_path.replace("\\", "/").strip())
    joined = posixpath.normpath(posixpath.join(md_dir, cleaned) if md_dir else cleaned)
    if joined.startswith("../") or joined == ".." or joined.startswith("/"):
        return None
    return joined


def convert_markdown_to_docx(
    markdown: str,
    *,
    images: Mapping[str, bytes | None] | None = None,
    layout: DocLayout = LAYOUT_STANDARD,
) -> MdToDocxResult:
    """Convert Markdown text to a .docx.

    ``images`` maps the raw ``src`` from the Markdown to image bytes, or ``None``
    when the file was looked up and missing. Remote / non-relative srcs are
    warned and rendered as alt text (+ URL when present).

    ``layout`` 选排版档位：``standard``（默认）= 技术文档/报告；``official`` =
    中文正式文书（首行缩进、两端对齐、公文页边距、页码）。
    """
    _ensure_docx()
    image_map = dict(images or {})
    warnings: list[str] = []
    indent_body = layout == LAYOUT_OFFICIAL
    doc = _DocumentFactory()
    _apply_document_defaults(doc, layout)

    tokens = _MD.parse(markdown or "")
    i = 0
    while i < len(tokens):
        token = tokens[i]
        t = token.type

        if t == "heading_open":
            level = int(token.tag[1]) if token.tag and token.tag.startswith("h") else 1
            level = max(1, min(level, 4))
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            p = doc.add_heading("", level=level)
            _style_heading_paragraph(p, level)
            if inline is not None and inline.type == "inline":
                _render_inline(p, inline, images=image_map, warnings=warnings, in_heading=True)
            i += 3  # open, inline, close
            continue

        if t == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            # Lone image paragraph → picture block (cleaner than inline-in-para).
            if (
                inline is not None
                and inline.type == "inline"
                and inline.children
                and len(inline.children) == 1
                and inline.children[0].type == "image"
            ):
                _render_image_block(
                    doc, inline.children[0], images=image_map, warnings=warnings
                )
            else:
                p = doc.add_paragraph()
                _style_body_paragraph(
                    p, first_line_indent=indent_body, justify=indent_body
                )
                if inline is not None and inline.type == "inline":
                    _render_inline(p, inline, images=image_map, warnings=warnings)
            i += 3
            continue

        if t == "bullet_list_open":
            i = _render_list(doc, tokens, i, ordered=False, images=image_map, warnings=warnings)
            continue

        if t == "ordered_list_open":
            i = _render_list(doc, tokens, i, ordered=True, images=image_map, warnings=warnings)
            continue

        if t == "fence":
            _render_fence(doc, token)
            i += 1
            continue

        if t == "code_block":
            _render_fence(doc, token)
            i += 1
            continue

        if t == "table_open":
            i = _render_table(doc, tokens, i, images=image_map, warnings=warnings)
            continue

        if t == "hr":
            doc.add_paragraph("─" * 24)
            i += 1
            continue

        if t == "blockquote_open":
            i = _render_blockquote(doc, tokens, i, images=image_map, warnings=warnings)
            continue

        # Skip structural closers / unknown opens we don't specially handle.
        i += 1

    # Warn about declared-missing images that never appeared as tokens (defensive).
    for src, data in image_map.items():
        if data is None:
            msg = f"缺图：{src}"
            if msg not in warnings:
                warnings.append(msg)

    buf = io.BytesIO()
    doc.save(buf)
    return MdToDocxResult(docx_bytes=buf.getvalue(), warnings=warnings)


# ---------------------------------------------------------------------------
# Document / run styling
# ---------------------------------------------------------------------------


def _apply_document_defaults(doc: Document, layout: DocLayout) -> None:
    section = doc.sections[0]
    section.page_width = _Cm(A4_WIDTH_CM)
    section.page_height = _Cm(A4_HEIGHT_CM)
    top, bottom, left, right = (
        OFFICIAL_MARGINS_CM if layout == LAYOUT_OFFICIAL else STANDARD_MARGINS_CM
    )
    section.top_margin = _Cm(top)
    section.bottom_margin = _Cm(bottom)
    section.left_margin = _Cm(left)
    section.right_margin = _Cm(right)
    if layout == LAYOUT_OFFICIAL:
        _add_page_number_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = _FONT_LATIN
    normal.font.size = _Pt(_BODY_PT)
    if normal._element.rPr is not None and normal._element.rPr.rFonts is not None:
        normal._element.rPr.rFonts.set(_qn("w:eastAsia"), _FONT_BODY_CJK)

    for level in range(1, 5):
        style = doc.styles[f"Heading {level}"]
        style.font.name = _FONT_LATIN
        style.font.size = _Pt(_HEADING_PT[level])
        style.font.bold = True
        style.font.color.rgb = _RGBColor(0x1F, 0x23, 0x28)
        if style._element.rPr is not None:
            r_fonts = style._element.rPr.rFonts
            if r_fonts is None:
                r_fonts = _OxmlElement("w:rFonts")
                style._element.rPr.append(r_fonts)
            r_fonts.set(_qn("w:eastAsia"), _FONT_HEADING_CJK)


def _set_run_font(
    run: Any,
    *,
    cjk: str,
    latin: str = _FONT_LATIN,
    size_pt: int | None = None,
) -> None:
    run.font.name = latin
    r = run._element
    r_pr = r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(_qn("w:eastAsia"), cjk)
    r_fonts.set(_qn("w:ascii"), latin)
    r_fonts.set(_qn("w:hAnsi"), latin)
    if size_pt is not None:
        run.font.size = _Pt(size_pt)


def _style_body_paragraph(
    p: Any, *, first_line_indent: bool = False, justify: bool = False
) -> None:
    p.paragraph_format.space_after = _Pt(BODY_SPACE_AFTER_PT)
    p.paragraph_format.line_spacing = LINE_SPACING_MULTIPLE
    if justify:
        p.alignment = _WD_ALIGN_PARAGRAPH.JUSTIFY
    if first_line_indent:
        _apply_first_line_indent(p)


def _style_list_paragraph(p: Any) -> None:
    p.paragraph_format.space_after = _Pt(BODY_SPACE_AFTER_PT)
    p.paragraph_format.line_spacing = LINE_SPACING_MULTIPLE


def _add_page_number_footer(section: Any) -> None:
    """official 档：页码底端居中（PAGE 域）。standard 不建页脚。"""
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = _WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    _set_run_font(run, cjk=_FONT_BODY_CJK, size_pt=9)
    r = run._r
    begin = _OxmlElement("w:fldChar")
    begin.set(_qn("w:fldCharType"), "begin")
    instr = _OxmlElement("w:instrText")
    instr.set(_qn("xml:space"), "preserve")
    instr.text = " PAGE "
    sep = _OxmlElement("w:fldChar")
    sep.set(_qn("w:fldCharType"), "separate")
    cached = _OxmlElement("w:t")
    cached.text = "1"
    end = _OxmlElement("w:fldChar")
    end.set(_qn("w:fldCharType"), "end")
    r.append(begin)
    r.append(instr)
    r.append(sep)
    r.append(cached)
    r.append(end)


def _apply_first_line_indent(p: Any) -> None:
    """首行缩进两字：``w:firstLineChars`` 为准，``w:firstLine`` 是给不认它的阅读器的兜底。

    Word 的「缩进 2 字符」真身是 ``firstLineChars``（百分之一字符为单位，随字号走）；
    只写 twips 的话换字号就不再是整两字。两个属性一起写 = Word 自己的写法。
    """
    p.paragraph_format.first_line_indent = _Pt(_BODY_PT * FIRST_LINE_INDENT_CHARS)
    ind = p._element.get_or_add_pPr().get_or_add_ind()
    ind.set(_qn("w:firstLineChars"), str(FIRST_LINE_INDENT_CHARS * 100))


def _style_heading_paragraph(p: Any, level: int) -> None:
    p.paragraph_format.space_before = _Pt(12 if level <= 2 else 8)
    p.paragraph_format.space_after = _Pt(6)
    if level == 1:
        # 文档大标题居中是 Word 通例（公文与技术报告都成立），故两档默认都开。
        p.alignment = _WD_ALIGN_PARAGRAPH.CENTER


# ---------------------------------------------------------------------------
# Block renderers
# ---------------------------------------------------------------------------


def _render_fence(doc: Document, token: Token) -> None:
    text = token.content or ""
    if text.endswith("\n"):
        text = text[:-1]
    lang = (token.info or "").strip()
    if lang:
        label = doc.add_paragraph()
        run = label.add_run(lang)
        _set_run_font(run, cjk=_FONT_BODY_CJK, latin=_FONT_CODE, size_pt=9)
        run.font.color.rgb = _RGBColor(0x6B, 0x72, 0x80)
        label.paragraph_format.space_after = _Pt(0)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = _Pt(8)
    run = p.add_run(text)
    _set_run_font(run, cjk=_FONT_BODY_CJK, latin=_FONT_CODE, size_pt=_CODE_PT)
    # Light shading via paragraph shading
    p_pr = p._element.get_or_add_pPr()
    shd = _OxmlElement("w:shd")
    shd.set(_qn("w:fill"), "F4F4F5")
    shd.set(_qn("w:val"), "clear")
    p_pr.append(shd)


def _render_list(
    doc: Document,
    tokens: list[Token],
    start: int,
    *,
    ordered: bool,
    images: dict[str, bytes | None],
    warnings: list[str],
    level: int = 0,
) -> int:
    i = start + 1
    close_type = "ordered_list_close" if ordered else "bullet_list_close"
    style = "List Number" if ordered else "List Bullet"
    while i < len(tokens):
        token = tokens[i]
        if token.type == close_type:
            return i + 1
        if token.type == "list_item_open":
            i += 1
            # Consume item body until list_item_close.
            while i < len(tokens) and tokens[i].type != "list_item_close":
                if tokens[i].type == "paragraph_open":
                    inline = tokens[i + 1] if i + 1 < len(tokens) else None
                    p = doc.add_paragraph(style=style)
                    _style_list_paragraph(p)
                    if level:
                        p.paragraph_format.left_indent = _Cm(0.75 * level)
                    if inline is not None and inline.type == "inline":
                        _render_inline(p, inline, images=images, warnings=warnings)
                    i += 3
                    continue
                if tokens[i].type in ("bullet_list_open", "ordered_list_open"):
                    nested_ordered = tokens[i].type == "ordered_list_open"
                    i = _render_list(
                        doc,
                        tokens,
                        i,
                        ordered=nested_ordered,
                        images=images,
                        warnings=warnings,
                        level=level + 1,
                    )
                    continue
                i += 1
            if i < len(tokens) and tokens[i].type == "list_item_close":
                i += 1
            continue
        i += 1
    return i


def _render_table(
    doc: Document,
    tokens: list[Token],
    start: int,
    *,
    images: dict[str, bytes | None],
    warnings: list[str],
) -> int:
    # Collect rows of cell inlines.
    rows: list[list[Token | None]] = []
    i = start + 1
    current: list[Token | None] = []
    while i < len(tokens):
        t = tokens[i]
        if t.type == "table_close":
            i += 1
            break
        if t.type == "tr_open":
            current = []
            i += 1
            continue
        if t.type == "tr_close":
            rows.append(current)
            current = []
            i += 1
            continue
        if t.type in ("th_open", "td_open"):
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            current.append(inline if inline is not None and inline.type == "inline" else None)
            # skip open, inline, close
            i += 3
            continue
        i += 1

    if not rows:
        return i
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            inline = row[c_idx] if c_idx < len(row) else None
            if inline is not None:
                _render_inline(p, inline, images=images, warnings=warnings)
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()
    return i


def _render_blockquote(
    doc: Document,
    tokens: list[Token],
    start: int,
    *,
    images: dict[str, bytes | None],
    warnings: list[str],
) -> int:
    i = start + 1
    while i < len(tokens):
        if tokens[i].type == "blockquote_close":
            return i + 1
        if tokens[i].type == "paragraph_open":
            inline = tokens[i + 1] if i + 1 < len(tokens) else None
            p = doc.add_paragraph()
            # 引用块不吃 official 档的首行缩进：整块左缩进 + 竖线已经把它和正文分开了。
            _style_body_paragraph(p)
            p.paragraph_format.left_indent = _Cm(0.75)
            run_prefix = p.add_run("｜ ")
            _set_run_font(run_prefix, cjk=_FONT_BODY_CJK, size_pt=_BODY_PT)
            run_prefix.font.color.rgb = _RGBColor(0x9C, 0xA3, 0xAF)
            if inline is not None and inline.type == "inline":
                _render_inline(p, inline, images=images, warnings=warnings)
            i += 3
            continue
        i += 1
    return i


def _render_image_block(
    doc: Document,
    image_token: Token,
    *,
    images: dict[str, bytes | None],
    warnings: list[str],
) -> None:
    src = str(image_token.attrGet("src") or "").strip()
    alt = (image_token.content or "").strip() or str(image_token.attrGet("alt") or "").strip()
    data = images.get(src) if src in images else None
    if src in images and data is None:
        msg = f"缺图：{src}"
        if msg not in warnings:
            warnings.append(msg)
        p = doc.add_paragraph()
        run = p.add_run(f"[缺图：{alt or src}]")
        _set_run_font(run, cjk=_FONT_BODY_CJK, size_pt=_BODY_PT)
        run.font.color.rgb = _RGBColor(0xB9, 0x1C, 0x1C)
        return
    if not is_embeddable_relative_src(src) or data is None:
        if is_embeddable_relative_src(src) and src not in images:
            msg = f"缺图：{src}"
            if msg not in warnings:
                warnings.append(msg)
        elif not is_embeddable_relative_src(src):
            msg = f"跳过非相对路径图片：{src}"
            if msg not in warnings:
                warnings.append(msg)
        p = doc.add_paragraph()
        label = alt or src
        run = p.add_run(f"[图片：{label}]")
        _set_run_font(run, cjk=_FONT_BODY_CJK, size_pt=_BODY_PT)
        return
    try:
        doc.add_picture(io.BytesIO(data), width=_MAX_IMAGE_WIDTH)
    except Exception as exc:  # noqa: BLE001 — bad image bytes should warn, not abort
        msg = f"图片无法嵌入（{src}）：{exc}"
        if msg not in warnings:
            warnings.append(msg)
        p = doc.add_paragraph()
        run = p.add_run(f"[图片损坏：{alt or src}]")
        _set_run_font(run, cjk=_FONT_BODY_CJK, size_pt=_BODY_PT)
        run.font.color.rgb = _RGBColor(0xB9, 0x1C, 0x1C)
        return
    if alt:
        cap = doc.add_paragraph()
        cap.alignment = _WD_ALIGN_PARAGRAPH.CENTER
        run = cap.add_run(alt)
        _set_run_font(run, cjk=_FONT_BODY_CJK, size_pt=9)
        run.font.color.rgb = _RGBColor(0x6B, 0x72, 0x80)


# ---------------------------------------------------------------------------
# Inline renderers
# ---------------------------------------------------------------------------


def _render_inline(
    paragraph: Any,
    inline: Token,
    *,
    images: dict[str, bytes | None],
    warnings: list[str],
    in_heading: bool = False,
) -> None:
    children = inline.children or []
    stack_link: list[str] = []
    bold = 0
    italic = 0
    code = 0
    i = 0
    while i < len(children):
        child = children[i]
        ct = child.type
        if ct == "text":
            _add_styled_run(
                paragraph,
                child.content or "",
                bold=bold > 0,
                italic=italic > 0,
                code=code > 0,
                link_url=stack_link[-1] if stack_link else None,
                heading=in_heading,
            )
        elif ct == "code_inline":
            _add_styled_run(
                paragraph,
                child.content or "",
                bold=False,
                italic=False,
                code=True,
                link_url=stack_link[-1] if stack_link else None,
                heading=in_heading,
            )
        elif ct == "softbreak":
            _add_styled_run(paragraph, "\n", heading=in_heading)
        elif ct == "hardbreak":
            paragraph.add_run().add_break()
        elif ct == "strong_open":
            bold += 1
        elif ct == "strong_close":
            bold = max(0, bold - 1)
        elif ct == "em_open":
            italic += 1
        elif ct == "em_close":
            italic = max(0, italic - 1)
        elif ct == "link_open":
            href = str(child.attrGet("href") or "").strip()
            stack_link.append(href)
        elif ct == "link_close":
            if stack_link:
                stack_link.pop()
        elif ct == "image":
            # Inline image inside mixed paragraph — embed if possible, else alt.
            src = str(child.attrGet("src") or "").strip()
            alt = (child.content or "").strip() or str(child.attrGet("alt") or "").strip()
            data = images.get(src) if src in images else None
            if src in images and data is None:
                msg = f"缺图：{src}"
                if msg not in warnings:
                    warnings.append(msg)
                _add_styled_run(paragraph, f"[缺图：{alt or src}]", heading=in_heading)
            elif data:
                try:
                    run = paragraph.add_run()
                    run.add_picture(io.BytesIO(data), width=_Inches(3.2))
                except Exception as exc:  # noqa: BLE001
                    msg = f"图片无法嵌入（{src}）：{exc}"
                    if msg not in warnings:
                        warnings.append(msg)
                    _add_styled_run(paragraph, f"[图片损坏：{alt or src}]", heading=in_heading)
            else:
                if is_embeddable_relative_src(src):
                    msg = f"缺图：{src}"
                    if msg not in warnings:
                        warnings.append(msg)
                else:
                    msg = f"跳过非相对路径图片：{src}"
                    if msg not in warnings:
                        warnings.append(msg)
                _add_styled_run(paragraph, f"[图片：{alt or src}]", heading=in_heading)
        elif ct == "html_inline":
            # html=False should not emit these; ignore defensively.
            pass
        i += 1


def _add_styled_run(
    paragraph: Any,
    text: str,
    *,
    bold: bool = False,
    italic: bool = False,
    code: bool = False,
    link_url: str | None = None,
    heading: bool = False,
) -> None:
    if not text and not link_url:
        return
    if link_url and _is_safe_http_url(link_url):
        run = _add_hyperlink(paragraph, text or link_url, link_url)
    else:
        run = paragraph.add_run(text)
    run.bold = bold or heading
    run.italic = italic
    if code:
        _set_run_font(run, cjk=_FONT_BODY_CJK, latin=_FONT_CODE, size_pt=_CODE_PT)
        run.font.color.rgb = _RGBColor(0x37, 0x40, 0x51)
    elif heading:
        _set_run_font(run, cjk=_FONT_HEADING_CJK, size_pt=None)
    else:
        _set_run_font(run, cjk=_FONT_BODY_CJK, size_pt=_BODY_PT)
    if link_url and _is_safe_http_url(link_url):
        run.font.color.rgb = _RGBColor(0x05, 0x63, 0xC1)
        run.underline = True


def _is_safe_http_url(url: str) -> bool:
    try:
        parsed = urlparse(url)
    except ValueError:
        return False
    return parsed.scheme in ("http", "https") and bool(parsed.netloc)


def _add_hyperlink(paragraph: Any, text: str, url: str) -> Any:
    """Insert an external hyperlink run into ``paragraph`` (python-docx has no helper)."""
    part = paragraph.part
    r_id = part.relate_to(url, _RT.HYPERLINK, is_external=True)
    hyperlink = _OxmlElement("w:hyperlink")
    hyperlink.set(_qn("r:id"), r_id)
    new_run = _OxmlElement("w:r")
    r_pr = _OxmlElement("w:rPr")
    new_run.append(r_pr)
    text_elem = _OxmlElement("w:t")
    text_elem.text = text
    new_run.append(text_elem)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    # Wrap as a Run-like object for font helpers.
    from docx.text.run import Run

    return Run(new_run, paragraph)
