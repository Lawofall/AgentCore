"""Tests for the file_write, file_delete and file_move tools (mutating file ops).

Hermetic: every test runs against a throwaway ``ServerWorkspace`` rooted at
``tmp_path`` and inspects the real on-disk result, mirroring the str_replace tool
tests. These tools are thin shells, so the focus is argument handling and the
typed-failure → user-message mapping (the heavy I/O lives in the backend).
"""

from dataclasses import replace
from io import BytesIO
from pathlib import Path
from unittest.mock import AsyncMock, patch

import pytest

from agentcore.tools.builtin.file_ops import (
    FileAppendTool,
    FileBatchTool,
    FileCopyTool,
    FileDeleteTool,
    FileListTool,
    FileMoveTool,
    FileReadTool,
    FileWriteTool,
    GlobTool,
    MkdirTool,
    StrReplaceTool,
    expand_brace_globs,
)
from agentcore.tools.builtin.file_ops.listing import (
    GLOB_DEPTH,
    GlobPlan,
    compile_glob_pattern,
    compile_glob_patterns,
)
from agentcore.tools.protocol import ToolContext
from agentcore.tools.sandbox.subprocess import SubprocessSandbox
from agentcore.workspace.server import ServerWorkspace
from agentcore.workspace.stage_dirs import REVIEWS_PREFIX


def _ctx(workspace: Path, *, agent_id: str = "a") -> ToolContext:
    # These tests are not empty-desk cases. A visible user file keeps project-shell
    # strip from rewriting nested paths the assertions pin.
    keep = workspace / "README.md"
    if not keep.exists():
        keep.write_text("desk\n", encoding="utf-8")
    return ToolContext.create(
        execution_id="e",
        run_id="s",
        agent_id=agent_id,
        backend=ServerWorkspace(root=workspace, sandbox=SubprocessSandbox()),
        user_id="u",
    )


def _docx_bytes(*, paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    from docx import Document

    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if table:
        grid = doc.add_table(rows=len(table), cols=max(len(row) for row in table))
        for i, row in enumerate(table):
            for j, value in enumerate(row):
                grid.cell(i, j).text = value
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _inprocess_extract():
    from agentcore.workspace.attachment_parse import (
        extract_office_from_disk,
        extract_office_payload,
    )

    def _run(data, ext, timeout, *, holder=None, source_path=None, start_page=1):
        if source_path is not None:
            return extract_office_from_disk(
                Path(source_path), ext, start_page=start_page
            )
        return extract_office_payload(data or b"", ext, start_page=start_page)

    return patch(
        "agentcore.workspace.attachment_parse._run_extract_subprocess",
        side_effect=_run,
    )


# --- file_write ---


async def test_write_creates_file(tmp_path: Path):
    result = await FileWriteTool().execute(
        {"path": "notes/report.md", "content": "# Hi"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "notes" / "report.md").read_text(encoding="utf-8") == "# Hi"


async def test_write_rejects_cleared_stub_content(tmp_path: Path):
    """过渡硬拒：content=[已清理] / 已落盘摘要不得写盘；正常长文仍成功。"""
    target = tmp_path / "doc.md"
    target.write_text("keep-me", encoding="utf-8")

    stub = await FileWriteTool().execute(
        {"path": "doc.md", "content": "[已清理]"}, _ctx(tmp_path)
    )
    assert stub.success is False
    assert stub.contract_failure is True
    assert "不能写入磁盘" in (stub.error or "")
    assert target.read_text(encoding="utf-8") == "keep-me"

    landed = await FileWriteTool().execute(
        {
            "path": "doc.md",
            "_landed_summary": "【已落盘摘要·只读】file_write 已成功写入",
            "status": "landed",
        },
        _ctx(tmp_path),
    )
    assert landed.success is False
    assert landed.contract_failure is True
    assert target.read_text(encoding="utf-8") == "keep-me"

    # Prose mentioning 已清理 must not be blocked.
    prose = "本节已清理历史遗留问题。" + ("正文。" * 200)
    ok = await FileWriteTool().execute({"path": "doc.md", "content": prose}, _ctx(tmp_path))
    assert ok.success is True
    assert target.read_text(encoding="utf-8") == prose


async def test_str_replace_rejects_cleared_stub_new_string(tmp_path: Path):
    (tmp_path / "notes.md").write_text("alpha\nbeta\n", encoding="utf-8")
    result = await StrReplaceTool().execute(
        {
            "path": "notes.md",
            "old_string": "alpha",
            "new_string": "[已清理·须重填]",
        },
        _ctx(tmp_path),
    )
    assert result.success is False
    assert result.contract_failure is True
    assert (tmp_path / "notes.md").read_text(encoding="utf-8") == "alpha\nbeta\n"


async def test_write_allows_substantial_overwrite(tmp_path: Path):
    """成篇 md/html 整盖允许（不再硬拒）；磁盘被新内容覆盖。"""
    body = "成篇正文。" * 80  # well over substantial threshold
    target = tmp_path / "site" / "index.html"
    target.parent.mkdir(parents=True)
    target.write_text(body, encoding="utf-8")
    new_body = "<html><body>rewrite complete page</body></html>"
    result = await FileWriteTool().execute(
        {"path": "site/index.html", "content": new_body},
        _ctx(tmp_path),
    )
    assert result.success is True
    assert target.read_text(encoding="utf-8") == new_body


async def test_write_allows_tiny_overwrite(tmp_path: Path):
    (tmp_path / "stub.txt").write_text("tiny", encoding="utf-8")
    result = await FileWriteTool().execute(
        {"path": "stub.txt", "content": "still small"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "stub.txt").read_text(encoding="utf-8") == "still small"


async def test_write_allows_non_empty_code_overwrite(tmp_path: Path):
    """非空代码整盖允许（优先局部劝导；不再硬拒）。"""
    body = "export function TopBar() {\n  return <header>App</header>;\n}\n"
    target = tmp_path / "src" / "TopBar.tsx"
    target.parent.mkdir(parents=True)
    target.write_text(body, encoding="utf-8")
    rewritten = "export function TopBar() {\n  return <header>New</header>;\n}\n"
    result = await FileWriteTool().execute(
        {"path": "src/TopBar.tsx", "content": rewritten},
        _ctx(tmp_path),
    )
    assert result.success is True
    assert target.read_text(encoding="utf-8") == rewritten


async def test_write_allows_css_js_overwrite(tmp_path: Path):
    """非空 css/js 整盖成功。"""
    css_old = "body { color: red; }\n" + "/* pad */\n" * 20
    js_old = "export const x = 1;\n" + "// pad\n" * 20
    (tmp_path / "styles").mkdir()
    (tmp_path / "scripts").mkdir()
    (tmp_path / "styles" / "main.css").write_text(css_old, encoding="utf-8")
    (tmp_path / "scripts" / "app.js").write_text(js_old, encoding="utf-8")
    css_new = "body { color: blue; }\n"
    js_new = "export const x = 2;\n"
    css_r = await FileWriteTool().execute(
        {"path": "styles/main.css", "content": css_new}, _ctx(tmp_path)
    )
    js_r = await FileWriteTool().execute(
        {"path": "scripts/app.js", "content": js_new}, _ctx(tmp_path)
    )
    assert css_r.success is True
    assert js_r.success is True
    assert (tmp_path / "styles" / "main.css").read_text(encoding="utf-8") == css_new
    assert (tmp_path / "scripts" / "app.js").read_text(encoding="utf-8") == js_new


async def test_write_allows_empty_code_shell(tmp_path: Path):
    """真·空壳（空白）代码文件仍可用 file_write 写入。"""
    target = tmp_path / "src" / "NewWidget.tsx"
    target.parent.mkdir(parents=True)
    target.write_text("   \n", encoding="utf-8")
    content = "export function NewWidget() {\n  return null;\n}\n"
    result = await FileWriteTool().execute(
        {"path": "src/NewWidget.tsx", "content": content},
        _ctx(tmp_path),
    )
    assert result.success is True
    assert target.read_text(encoding="utf-8") == content


async def test_write_allows_new_code_file(tmp_path: Path):
    content = "export const x = 1;\n"
    result = await FileWriteTool().execute(
        {"path": "src/fresh.ts", "content": content},
        _ctx(tmp_path),
    )
    assert result.success is True
    assert (tmp_path / "src" / "fresh.ts").read_text(encoding="utf-8") == content


async def test_write_rejects_empty_path(tmp_path: Path):
    # A worker that omits/empties ``path`` must get a crisp required-arg error — NOT
    # a backend write onto the workspace root dir (the real-world file_write failure:
    # path=None → root → "[Errno 13] Permission denied: <abs server path>").
    (tmp_path / "keep.txt").write_text("keep", encoding="utf-8")
    result = await FileWriteTool().execute({"path": "", "content": "x" * 5000}, _ctx(tmp_path))
    assert result.success is False
    assert "path 不能为空" in result.error
    # the root must be untouched (no clobber, no stray file)
    assert (tmp_path / "keep.txt").read_text(encoding="utf-8") == "keep"


async def test_write_rejects_missing_path(tmp_path: Path):
    result = await FileWriteTool().execute({"content": "body"}, _ctx(tmp_path))
    assert result.success is False
    assert "path 不能为空" in result.error


async def test_write_rejects_path_outside_workspace(tmp_path: Path):
    ws = tmp_path / "ws"
    ws.mkdir()
    result = await FileWriteTool().execute({"path": "../escaped.md", "content": "leak"}, _ctx(ws))
    assert result.success is False
    assert "超出了工作区范围" in result.error
    assert not (tmp_path / "escaped.md").exists()


async def test_write_normalizes_absolute_workspace_path(tmp_path: Path):
    # A worker passing an absolute /workspace/... path now succeeds (normalized at the
    # path-resolution seam) instead of failing OutsideWorkspace and retrying.
    result = await FileWriteTool().execute(
        {"path": "/workspace/research/x.md", "content": "hi"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "research" / "x.md").read_text(encoding="utf-8") == "hi"


async def test_outside_workspace_error_is_actionable(tmp_path: Path):
    # The rejection tells the model exactly how to fix it (relative path + example),
    # not just that the path was out of range. Cloud workspace also nudges the bind card.
    ws = tmp_path / "ws"
    ws.mkdir()
    result = await FileWriteTool().execute({"path": "../escaped.md", "content": "x"}, _ctx(ws))
    assert result.success is False
    assert "超出了工作区范围" in result.error
    assert result.failure_code == "outside_workspace"
    assert "相对路径" in result.error
    assert "AgentCore/文档/research/report.md" in result.error
    assert "bind_local_folder" in result.error or "open_local_project" in result.error
    assert "open_local_project" in result.error or "本机传统" in result.error
    assert "导入到云" in result.error
    assert "≠离线" in result.error


# --- file_read ---


async def test_file_read_docx_transparent_extract(tmp_path: Path):
    (tmp_path / "docs").mkdir()
    (tmp_path / "docs" / "brief.docx").write_bytes(
        _docx_bytes(
            paragraphs=["Hello from docx with enough alphanumeric body for scan."],
            table=[["Name", "Qty"], ["Widget", "42"]],
        )
    )
    ctx = _ctx(tmp_path)
    with patch(
        "agentcore.workspace.attachment_parse._convert_with_markitdown",
        side_effect=AssertionError("docx must not call markitdown"),
    ):
        result = await FileReadTool().execute({"path": "docs/brief.docx"}, ctx)
    assert result.success is True
    assert "Hello from docx" in (result.output or "")
    assert "Widget" in (result.output or "")
    assert not (tmp_path / "docs" / "brief.docx.md").exists()


def test_sniff_bytes_magic():
    from agentcore.workspace.file_kind import sniff_bytes

    assert sniff_bytes(b"\xd0\xcf\x11\xe0xxxx") == "ole"
    assert sniff_bytes(b"%PDF-1.4") == "pdf"
    assert sniff_bytes(b"PK\x03\x04") == "zip"
    assert sniff_bytes(b"hello\n") == "text"


def test_extract_failed_text_hides_library_tokens():
    from agentcore.tools.builtin.file_ops.observe import extract_failed_text

    timeout = extract_failed_text("extract_timeout")
    assert "超时" in timeout
    assert "extract_timeout" not in timeout
    failed = extract_failed_text("convert:PackageNotFoundError")
    assert "抽文本失败" in failed
    assert "PackageNotFound" not in failed
    assert "markitdown" not in failed.lower()


async def test_file_read_pdf_transparent_extract(tmp_path: Path):
    (tmp_path / "paper.pdf").write_bytes(b"%PDF-fake")
    body = "Abstract\n\nThis paper studies agents." + ("x" * 20)
    with _inprocess_extract(), patch(
        "agentcore.workspace.attachment_parse._extract_pdf_text",
        return_value=(body, "markitdown"),
    ):
        result = await FileReadTool().execute({"path": "paper.pdf"}, _ctx(tmp_path))
    assert result.success is True
    assert "This paper studies agents" in (result.output or "")
    assert not (tmp_path / "paper.pdf.md").exists()


async def test_file_read_pdf_start_page_skips_sidecar(tmp_path: Path):
    from agentcore.workspace.attachment_parse import ExtractResult, ParseStatus

    (tmp_path / "paper.pdf").write_bytes(b"%PDF-x")
    (tmp_path / "paper.pdf.md").write_text(
        "Sidecar first window only with alphanumeric body.\n", encoding="utf-8"
    )
    with patch(
        "agentcore.tools.builtin.file_ops.read._extract_office",
        new=AsyncMock(
            return_value=ExtractResult(
                status=ParseStatus.OK,
                text="Clause from page 41 with enough alphanumeric body here.",
                detail="ok:pdfminer",
                page_start=41,
                page_end=80,
                page_total=200,
            )
        ),
    ) as mocked:
        result = await FileReadTool().execute(
            {"path": "paper.pdf", "start_page": 41}, _ctx(tmp_path)
        )
    assert mocked.await_args is not None
    assert mocked.await_args.kwargs.get("start_page") == 41
    assert result.success is True
    out = result.output or ""
    assert "Clause from page 41" in out
    assert "start_page=81" in out
    assert "Sidecar first window" not in out

async def test_file_read_xlsx_does_not_extract(tmp_path: Path):
    from unittest.mock import patch

    (tmp_path / "report.xlsx").write_bytes(b"PK\x03\x04")
    with patch(
        "agentcore.tools.builtin.file_ops.read._code_execute_assembled",
        return_value=True,
    ):
        result = await FileReadTool().execute({"path": "report.xlsx"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert "用 run" in out
    assert "code_execute" not in out
    assert "[观察信封]" in out
    assert "kind: table" in out
    assert not (tmp_path / "report.xlsx.md").exists()


async def test_file_read_table_without_code_execute_omits_tool_name(tmp_path: Path):
    from unittest.mock import patch

    (tmp_path / "report.xlsx").write_bytes(b"PK\x03\x04")
    (tmp_path / "upload.csv").write_text("a,b\n1,2\n", encoding="utf-8")
    with patch(
        "agentcore.tools.builtin.file_ops.read._code_execute_assembled",
        return_value=False,
    ):
        xlsx = await FileReadTool().execute({"path": "report.xlsx"}, _ctx(tmp_path))
        csv = await FileReadTool().execute({"path": "upload.csv"}, _ctx(tmp_path))
    assert xlsx.success is True
    assert csv.success is True
    assert "code_execute" not in (xlsx.output or "")
    assert "请用 run" not in (xlsx.output or "")
    assert "code_execute" not in (csv.output or "")
    assert "手抄" in (xlsx.output or "")
    assert "结构报告" in (xlsx.output or "")
    assert "待跑" in (xlsx.output or "")
    assert "无法可靠处理" not in (xlsx.output or "")


async def test_file_read_landed_csv_is_readable(tmp_path: Path):
    """Worker 自产表格认落盘台账，可 file_read 回读；未入账的同扩展名仍拒。"""
    ctx = _ctx(tmp_path)
    written = await FileWriteTool().execute(
        {"path": "out/data.csv", "content": "name,n\nalice,1\n"},
        ctx,
    )
    assert written.success is True
    assert ctx.landed_artifact_kinds.get("out/data.csv") is not None

    ok = await FileReadTool().execute({"path": "out/data.csv"}, ctx)
    assert ok.success is True
    assert "alice,1" in (ok.output or "")

    (tmp_path / "foreign.csv").write_text("x,y\n9,8\n", encoding="utf-8")
    blocked = await FileReadTool().execute({"path": "foreign.csv"}, ctx)
    assert blocked.success is True
    assert "code_execute" not in (blocked.output or "")
    assert "kind: table" in (blocked.output or "")


async def test_file_read_scanned_pdf_notice(tmp_path: Path):
    (tmp_path / "scan.pdf").write_bytes(b"%PDF" + b"\x00" * 100)
    with _inprocess_extract(), patch(
        "agentcore.workspace.attachment_parse._extract_pdf_text",
        return_value=("   \n", "pdfminer"),
    ):
        result = await FileReadTool().execute({"path": "scan.pdf"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert "扫描" in out or "OCR" in out
    assert "[观察信封]" in out
    assert "kind: scan" in out
    assert "请用户" not in out
    assert not (tmp_path / "scan.pdf.md").exists()


async def test_file_read_legacy_doc_is_ole_envelope(tmp_path: Path):
    (tmp_path / "memo.doc").write_bytes(b"\xd0\xcf\x11\xe0" + b"\x00" * 32)
    result = await FileReadTool().execute({"path": "memo.doc"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert "[观察信封]" in out
    assert "ole" in out.lower()
    assert "codec" not in out.lower()
    assert "请用户" not in out


async def test_file_read_utf8_fail_pdf_magic_routes_to_extract(tmp_path: Path):
    from unittest.mock import AsyncMock, patch

    from agentcore.workspace.attachment_parse import ExtractResult, ParseStatus

    (tmp_path / "notes.txt").write_bytes(b"%PDF-" + b"\x00" * 40)
    with patch(
        "agentcore.tools.builtin.file_ops.read._extract_office",
        new=AsyncMock(
            return_value=ExtractResult(
                status=ParseStatus.OK, text="Abstract from sniffed PDF\n", detail="ok"
            )
        ),
    ):
        result = await FileReadTool().execute({"path": "notes.txt"}, _ctx(tmp_path))
    assert result.success is True
    assert "Abstract from sniffed PDF" in (result.output or "")


async def test_file_read_office_offset_limit_on_extracted_lines(tmp_path: Path):
    from agentcore.workspace.attachment_parse import ExtractResult, ParseStatus

    body = "\n".join(f"line-{i}" for i in range(1, 11))
    body = body + "\n" + ("word " * 20)
    (tmp_path / "notes.docx").write_bytes(b"PK")
    ctx = _ctx(tmp_path)
    with patch(
        "agentcore.tools.builtin.file_ops.read._extract_office",
        new=AsyncMock(
            return_value=ExtractResult(status=ParseStatus.OK, text=body, detail="ok")
        ),
    ):
        result = await FileReadTool().execute(
            {"path": "notes.docx", "offset": 2, "limit": 3},
            ctx,
        )
    assert result.success is True
    out = result.output or ""
    assert "line-2" in out
    assert "line-4" in out
    assert "line-1" not in out
    assert "共 " in out


async def test_file_read_prefers_existing_md_sidecar(tmp_path: Path):
    (tmp_path / "memo.docx").write_bytes(b"PK-original")
    (tmp_path / "memo.docx.md").write_text(
        "Sidecar text already prepared with enough body.\n", encoding="utf-8"
    )
    with patch(
        "agentcore.tools.builtin.file_ops.read._extract_office",
        side_effect=AssertionError("extract must not run when sidecar exists"),
    ):
        result = await FileReadTool().execute({"path": "memo.docx"}, _ctx(tmp_path))
    assert result.success is True
    assert "Sidecar text already prepared" in (result.output or "")


async def test_file_read_scan_sidecar_is_scan_envelope(tmp_path: Path):
    from agentcore.workspace.attachment_parse import SCAN_NOTICE

    (tmp_path / "scan.pdf").write_bytes(b"%PDF-x")
    (tmp_path / "scan.pdf.md").write_text(SCAN_NOTICE + "\n", encoding="utf-8")
    with patch(
        "agentcore.tools.builtin.file_ops.read._extract_office",
        side_effect=AssertionError("extract must not run for scan sidecar"),
    ):
        result = await FileReadTool().execute({"path": "scan.pdf"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert "[观察信封]" in out
    assert "kind: scan" in out
    assert "请用户" not in out


async def test_file_read_office_extract_failure_soft(tmp_path: Path):
    (tmp_path / "broken.docx").write_bytes(b"not-a-docx")
    result = await FileReadTool().execute({"path": "broken.docx"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert "[观察信封]" in out
    assert "kind: extract" in out
    assert "抽文本失败" in out
    assert "convert:" not in out
    assert "markitdown" not in out.lower()
    assert "请用 code_execute" not in out
    assert "code_execute" not in out
    assert "read_image" in out
    assert "请用户" not in out

async def test_file_read_same_path_repeated_reads_return_disk_body(tmp_path: Path):
    """同一 path 连续多次成功读仍返回磁盘正文。"""
    (tmp_path / "doc.md").write_text("# Title\nunique-body-marker\n", encoding="utf-8")
    ctx = _ctx(tmp_path)
    tool = FileReadTool()
    for i in range(6):
        result = await tool.execute({"path": "doc.md"}, ctx)
        assert result.success is True, i
        out = result.output or ""
        assert "unique-body-marker" in out
        assert "Title" in out


async def test_file_read_pagination_windows_return_requested_lines(tmp_path: Path):
    lines = "\n".join(f"L{i}" for i in range(1, 81))
    (tmp_path / "page.md").write_text(lines + "\n", encoding="utf-8")
    ctx = _ctx(tmp_path)
    tool = FileReadTool()
    first = await tool.execute({"path": "page.md", "offset": 1, "limit": 5}, ctx)
    assert first.success is True
    assert "L1" in (first.output or "")
    assert "未达安全顶，省略 limit 可整读" in (first.output or "")
    assert "已达行顶" not in (first.output or "")
    expand = await tool.execute({"path": "page.md", "offset": 3, "limit": 8}, ctx)
    assert expand.success is True
    assert "L3" in (expand.output or "")
    for start in (11, 16, 21, 26, 31):
        result = await tool.execute(
            {"path": "page.md", "offset": start, "limit": 5}, ctx
        )
        assert result.success is True, start
        assert f"L{start}" in (result.output or "")


async def test_file_read_after_tool_clear_still_returns_disk_body(tmp_path: Path):
    """投影把旧结果收成指针后，再读仍给磁盘正文。"""
    import json

    from agentcore.llm.provider.protocol import LLMMessage, ToolCall, ToolCallFunction
    from agentcore.runtime.engine.tool_clear import project_cleared_window

    def pair(call_id: str, path: str, result: str) -> list[LLMMessage]:
        return [
            LLMMessage(
                role="assistant",
                content=None,
                tool_calls=[
                    ToolCall(
                        id=call_id,
                        function=ToolCallFunction(
                            name="file_read",
                            arguments=json.dumps({"path": path}),
                        ),
                    )
                ],
            ),
            LLMMessage(role="tool", content=result, tool_call_id=call_id),
        ]

    body = "\n".join(f"line-{i} " + ("x" * 80) for i in range(80))
    (tmp_path / "src").mkdir()
    (tmp_path / "src" / "target.py").write_text(body + "\n", encoding="utf-8")
    (tmp_path / "src" / "other.py").write_text(body + "\n", encoding="utf-8")
    ctx = _ctx(tmp_path)
    tool = FileReadTool()
    outputs: list[str] = []
    for _ in range(3):
        result = await tool.execute(
            {"path": "src/target.py", "offset": 1, "limit": 40}, ctx
        )
        assert result.success is True
        outputs.append(result.output or "")
        assert "line-0" in (result.output or "")
        assert "全文" not in (result.output or "")

    msgs: list[LLMMessage] = [LLMMessage(role="user", content="go")]
    for i, output in enumerate(outputs):
        msgs += pair(f"t{i}", "src/target.py", output)
    for i in range(2):
        other = await tool.execute({"path": "src/other.py"}, ctx)
        assert other.success is True
        msgs += pair(f"o{i}", "src/other.py", other.output or "")

    projected = project_cleared_window(
        msgs,
        clearable_tools=frozenset({"file_read"}),
        keep_recent=2,
        min_chars=100,
        summary_max_chars=0,
    )
    stub = next(m for m in projected if m.tool_call_id == "t0")
    assert (stub.content or "").startswith("[已清理")
    assert "path='src/target.py'" in (stub.content or "")
    assert "status=content_cleared" in (stub.content or "")
    assert "disk=intact" in (stub.content or "")
    assert "reread=omit_offset_limit" in (stub.content or "")

    recovered = await tool.execute({"path": "src/target.py"}, ctx)
    assert recovered.success is True
    assert "line-0" in (recovered.output or "")


async def test_file_write_then_file_read_returns_new_body(tmp_path: Path):
    """写成功后核对读仍返回磁盘新正文。"""
    (tmp_path / "draft.md").write_text(
        "\n".join(f"L{i}" for i in range(1, 12)) + "\n", encoding="utf-8"
    )
    ctx = _ctx(tmp_path)
    tool = FileReadTool()
    first = await tool.execute({"path": "draft.md", "offset": 2, "limit": 3}, ctx)
    assert first.success is True
    assert "L2" in (first.output or "")

    w = await FileWriteTool().execute(
        {"path": "draft.md", "content": "\n".join(f"N{i}" for i in range(1, 12)) + "\n"},
        ctx,
    )
    assert w.success is True
    assert ctx.landed_artifact_kinds.get("draft.md") is not None
    verify = await tool.execute({"path": "draft.md", "offset": 2, "limit": 3}, ctx)
    assert verify.success is True
    assert "N2" in (verify.output or "")


async def test_file_read_long_file_window_footer_and_output_limit(tmp_path: Path):
    """超行顶：编号窗 + 「共 N 行」+ 行顶页脚；无「中间省略」；output_limit 覆盖全文。"""
    from agentcore.tools.builtin.file_ops import FILE_READ_SAFETY_LINE_CAP

    total = FILE_READ_SAFETY_LINE_CAP + 50
    body = "\n".join(f"line-{i}" for i in range(1, total + 1)) + "\n"
    (tmp_path / "long.md").write_text(body, encoding="utf-8")
    result = await FileReadTool().execute({"path": "long.md"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert f"共 {total} 行" in out
    assert f"第 1–{FILE_READ_SAFETY_LINE_CAP} 行" in out
    assert "已达行顶" in out
    assert f"全文 {total} 行" not in out
    assert f"line-{FILE_READ_SAFETY_LINE_CAP}" in out
    assert f"line-{total}" not in out  # beyond window
    assert "中间省略" not in out
    assert "已保留首尾" not in out
    assert "系统视图截断" not in out
    assert result.output_limit is not None
    assert result.output_limit >= len(out)
    assert "     1|line-1" in out


async def test_file_read_800_lines_returns_full_text(tmp_path: Path):
    """800 行一次全文；未截断页脚「全文 N 行」。"""
    total = 800
    body = "\n".join(f"line-{i}" for i in range(1, total + 1)) + "\n"
    (tmp_path / "mid.md").write_text(body, encoding="utf-8")
    ctx = _ctx(tmp_path)
    result = await FileReadTool().execute({"path": "mid.md"}, ctx)
    assert result.success is True
    out = result.output or ""
    assert f"全文 {total} 行" in out
    assert "已达行顶" not in out
    assert "已达字符顶" not in out
    assert "     1|line-1" in out
    assert f"line-{total}" in out
    assert result.output_limit is not None
    assert result.output_limit >= len(out)


async def test_file_read_char_cap_truncates_complete_lines(tmp_path: Path):
    """超 8 万字符：先到先停、完整行；页脚标明字符顶。"""
    from agentcore.tools.builtin.file_ops import FILE_READ_SAFETY_CHAR_CAP

    line = "x" * 1000
    total = 90
    (tmp_path / "wide.md").write_text("\n".join([line] * total) + "\n", encoding="utf-8")
    result = await FileReadTool().execute({"path": "wide.md"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert "已达字符顶" in out
    assert f"共 {total} 行" in out
    assert "中间省略" not in out
    assert "已保留首尾" not in out
    # 1001*n - 1 > cap → n > 79.92, so 79 complete lines.
    expected_end = 0
    chars = 0
    for n in range(1, total + 1):
        extra = len(line) if n == 1 else 1 + len(line)
        if chars + extra > FILE_READ_SAFETY_CHAR_CAP:
            break
        chars += extra
        expected_end = n
    assert expected_end == 79
    assert f"第 1–{expected_end} 行" in out
    assert f"{expected_end:>6}|{line}" in out
    assert f"{expected_end + 1:>6}|{line}" not in out
    for raw in out.splitlines():
        if "|" in raw and raw.rsplit("|", 1)[-1].startswith("x"):
            assert raw.rsplit("|", 1)[-1] == line
    assert result.output_limit is not None
    assert result.output_limit >= len(out)


async def test_file_read_oversized_single_line_kept_whole(tmp_path: Path):
    """单行超字符顶：整行留下并标字符顶，禁止半行静默切。"""
    from agentcore.tools.builtin.file_ops import FILE_READ_SAFETY_CHAR_CAP

    line = "y" * (FILE_READ_SAFETY_CHAR_CAP + 50)
    (tmp_path / "one.md").write_text(line + "\n", encoding="utf-8")
    result = await FileReadTool().execute({"path": "one.md"}, _ctx(tmp_path))
    assert result.success is True
    out = result.output or ""
    assert f"     1|{line}" in out
    assert "已达字符顶" in out
    assert "中间省略" not in out
    assert result.output_limit is not None
    assert result.output_limit >= len(out)


async def test_file_read_oversized_message_does_not_teach_offset_limit(tmp_path: Path):
    from agentcore.workspace.limits import WORKSPACE_READ_MAX_BYTES

    (tmp_path / "huge.txt").write_bytes(b"a" * (WORKSPACE_READ_MAX_BYTES + 1))
    result = await FileReadTool().execute({"path": "huge.txt"}, _ctx(tmp_path))
    assert result.success is False
    err = (result.error or "").lower()
    assert "offset" not in err
    assert "limit" not in err
    assert "mib" in err
    assert "请用户" not in (result.error or "")
    assert result.failure_code == "too_large"


async def test_file_read_office_default_uses_same_full_window(tmp_path: Path):
    from agentcore.tools.builtin.file_ops import FILE_READ_SAFETY_LINE_CAP
    from agentcore.workspace.attachment_parse import ExtractResult, ParseStatus

    body = "\n".join(f"line-{i} word extra" for i in range(1, 801))
    (tmp_path / "notes.docx").write_bytes(b"PK")
    ctx = _ctx(tmp_path)
    with patch(
        "agentcore.tools.builtin.file_ops.read._extract_office",
        new=AsyncMock(
            return_value=ExtractResult(status=ParseStatus.OK, text=body, detail="ok")
        ),
    ):
        result = await FileReadTool().execute({"path": "notes.docx"}, ctx)
    assert result.success is True
    out = result.output or ""
    assert "全文 800 行" in out
    assert "line-800" in out

    over = "\n".join(
        f"line-{i} word extra" for i in range(1, FILE_READ_SAFETY_LINE_CAP + 51)
    )
    (tmp_path / "big.docx").write_bytes(b"PK")
    with patch(
        "agentcore.tools.builtin.file_ops.read._extract_office",
        new=AsyncMock(
            return_value=ExtractResult(status=ParseStatus.OK, text=over, detail="ok")
        ),
    ):
        truncated = await FileReadTool().execute({"path": "big.docx"}, ctx)
    tout = truncated.output or ""
    assert truncated.success is True
    assert f"第 1–{FILE_READ_SAFETY_LINE_CAP} 行" in tout
    assert "已达行顶" in tout
    assert f"line-{FILE_READ_SAFETY_LINE_CAP + 50}" not in tout

def test_file_read_schema_teaches_default_full_read():
    from agentcore.tools.builtin.file_ops import FILE_READ_SAFETY_LINE_CAP

    schema = FileReadTool().schema
    limit = schema.parameters["properties"]["limit"]
    assert limit["maximum"] == FILE_READ_SAFETY_LINE_CAP
    assert "省略则尽量整读" in limit["description"]
    assert "超安全顶截断" in limit["description"]
    desc = schema.description
    assert "grep" in desc or "code_search" in desc
    assert "glob" in desc
    assert "dump" in desc
    assert "web_fetch" in desc
    assert "web_fetch" in schema.parameters["properties"]["path"]["description"]
    assert "默认不抽文本" in schema.parameters["properties"]["path"]["description"]
    offset = schema.parameters["properties"]["offset"]
    assert "开窗" in offset["description"]
    assert "code_execute" not in desc
    assert "请用 code_execute。" not in desc
    assert "artifact manifest" not in desc
    assert "start_page" not in desc


@pytest.mark.parametrize(
    "path",
    [
        "https://example.com/docs/a.md",
        "http://example.com/x",
        "HTTPS://example.com/x",
    ],
)
async def test_file_read_http_url_reroutes_to_web_fetch(tmp_path: Path, path: str):
    result = await FileReadTool().execute({"path": path}, _ctx(tmp_path))
    err = result.error or ""
    assert result.success is False
    assert result.contract_failure is True
    assert result.metadata.get("code") == "url_not_workspace_path"
    assert result.failure_code == "url_not_workspace_path"
    assert result.metadata.get("cross_turn_retry") == "futile"
    assert "web_fetch" in err
    assert "不要把 URL 改写成路径再重试" in err


async def test_file_read_http_filename_is_not_a_url(tmp_path: Path):
    (tmp_path / "http_client.py").write_text("ok\n", encoding="utf-8")
    result = await FileReadTool().execute({"path": "http_client.py"}, _ctx(tmp_path))
    assert result.success is True
    assert "ok" in result.output
    assert result.metadata.get("code") != "url_not_workspace_path"


async def test_write_lands_substantial_prose_with_omission(tmp_path: Path):
    """省略套话不再硬拒；正文照常落盘。"""
    body = ("完整段落内容填充字。" * 50) + "\n……（中间省略，已保留首尾）……\n" + ("尾段续写。" * 30)
    assert len(body.strip()) >= 400
    result = await FileWriteTool().execute(
        {"path": "essay.md", "content": body}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "essay.md").read_text(encoding="utf-8") == body
    assert "产物疑似不完整" not in (result.output or "")


async def test_str_replace_then_file_read_returns_new_body(tmp_path: Path):
    (tmp_path / "edit.md").write_text("hello world\n", encoding="utf-8")
    ctx = _ctx(tmp_path)
    ok = await StrReplaceTool().execute(
        {"path": "edit.md", "old_string": "world", "new_string": "AgentCore"},
        ctx,
    )
    assert ok.success is True
    verify = await FileReadTool().execute({"path": "edit.md"}, ctx)
    assert verify.success is True
    assert "AgentCore" in (verify.output or "")


async def test_str_replace_failure_does_not_change_disk(tmp_path: Path):
    (tmp_path / "edit.md").write_text("hello world\n", encoding="utf-8")
    ctx = _ctx(tmp_path)
    fail = await StrReplaceTool().execute(
        {"path": "edit.md", "old_string": "no-such-token", "new_string": "x"},
        ctx,
    )
    assert fail.success is False
    assert (tmp_path / "edit.md").read_text(encoding="utf-8") == "hello world\n"
    still = await FileReadTool().execute({"path": "edit.md"}, ctx)
    assert still.success is True
    assert "hello world" in (still.output or "")


async def test_file_read_missing_does_not_trip_circuit_breaker(tmp_path: Path):
    """PathNotFound (env / wrong path) must not warn or disable file_read."""
    from agentcore.runtime.loop_controller import LoopController, ToolAttempt

    tool = FileReadTool()
    c = LoopController(tool_failure_warn=2, tool_failure_disable=3)
    for i in range(6):
        result = await tool.execute({"path": f"ghost/missing-{i}.md"}, _ctx(tmp_path))
        assert result.success is False
        assert result.contract_failure is True
        c.record(
            [
                ToolAttempt(
                    f"miss-{i}",
                    "file_read",
                    success=False,
                    error_summary=result.error or "",
                    contract_failure=result.contract_failure,
                    meta=dict(result.metadata or {}),
                )
            ]
        )
    cb = c.tool_circuit_breaker()
    assert cb.disabled == ()
    assert cb.warned == ()
    assert c.tool_failure_count("file_read") == 0


# --- file_append ---


async def test_append_creates_file_when_missing(tmp_path: Path):
    result = await FileAppendTool().execute(
        {"path": "draft.md", "content": "# Intro"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "draft.md").read_text(encoding="utf-8") == "# Intro"


async def test_append_adds_to_existing_file(tmp_path: Path):
    (tmp_path / "draft.md").write_text("# Intro", encoding="utf-8")
    result = await FileAppendTool().execute(
        {"path": "draft.md", "content": "\n\n## Section 2"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "draft.md").read_text(encoding="utf-8") == "# Intro\n\n## Section 2"


async def test_append_rejects_empty_path(tmp_path: Path):
    result = await FileAppendTool().execute({"path": "", "content": "x"}, _ctx(tmp_path))
    assert result.success is False
    assert "path 不能为空" in result.error


async def test_append_rejects_directory_target(tmp_path: Path):
    (tmp_path / "pkg").mkdir()
    result = await FileAppendTool().execute({"path": "pkg", "content": "x"}, _ctx(tmp_path))
    assert result.success is False
    assert "不是文件" in result.error


async def test_append_receipt_echoes_merged_tail(tmp_path: Path):
    # append 回执改为 artifact manifest（含 end_preview），免掉纯回读自检。
    (tmp_path / "draft.md").write_text("# Intro", encoding="utf-8")
    result = await FileAppendTool().execute(
        {"path": "draft.md", "content": "\n\n## Section 2"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "## Section 2" in result.output  # end_preview / title_tree
    assert "artifact manifest" in result.output
    assert "优先用 manifest 验真" in result.output


async def test_write_receipt_notes_persisted(tmp_path: Path):
    # file_write 回执 = artifact manifest；优先 manifest 验真（非身份硬闸）。
    result = await FileWriteTool().execute(
        {"path": "report.md", "content": "# Hi\n\n## A\n"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "artifact manifest" in result.output
    assert "content_sha256:" in result.output
    assert "title_tree:" in result.output
    assert "优先用 manifest 验真" in result.output
    assert "kind: skeleton" in result.output


async def test_write_receipt_reports_chars_not_bytes(tmp_path: Path):
    """规模按字符报：中文正文标成「字节」会让作者以为写少了、回读空转。"""
    body = "# 民事起诉状\n\n原告：昝雯，住青海省西宁市。\n"
    assert len(body.encode("utf-8")) != len(body)  # 中文下两口径必然分叉
    result = await FileWriteTool().execute(
        {"path": "诉状.md", "content": body}, _ctx(tmp_path)
    )
    assert result.success is True
    assert f"已写入 {len(body)} 字符到 诉状.md" in result.output
    assert f"chars: {len(body)}" in result.output
    assert "字节" not in result.output


async def test_write_prose_then_append_rejected(tmp_path: Path):
    """成篇 file_write 后同 path file_append 硬拒（Artifact-first）。"""
    prose = "# 报告\n\n" + ("这是实质正文段落。" * 50)  # well over substantial
    assert len(prose) >= 400
    ctx = _ctx(tmp_path)
    w = await FileWriteTool().execute({"path": "essay.md", "content": prose}, ctx)
    assert w.success is True
    assert "kind: prose" in w.output
    assert ctx.landed_artifact_kinds.get("essay.md") == "prose"
    blocked = await FileAppendTool().execute(
        {"path": "essay.md", "content": "\n\n## 续章\n更多。"}, ctx
    )
    assert blocked.success is False
    assert blocked.contract_failure is True
    assert "拒绝追加" in (blocked.error or "")
    assert "str_replace" in (blocked.error or "")
    assert "骨架填空" in (blocked.error or "") or "骨架" in (blocked.error or "")
    assert "应先短骨架" not in (blocked.error or "")
    assert "长交付物应先" not in (blocked.error or "")


async def test_write_skeleton_then_append_allowed(tmp_path: Path):
    skeleton = "# 报告\n\n## 一\n\n## 二\n\n<!-- OUTLINE -->\n"
    ctx = _ctx(tmp_path)
    w = await FileWriteTool().execute({"path": "report.md", "content": skeleton}, ctx)
    assert w.success is True
    assert ctx.landed_artifact_kinds.get("report.md") == "skeleton"
    a = await FileAppendTool().execute(
        {"path": "report.md", "content": "\n\n## 一\n\n正文填空。\n"}, ctx
    )
    assert a.success is True
    assert "artifact manifest" in a.output


async def test_file_read_allows_author_self_product(tmp_path: Path):
    """作者写后 body file_read 允许。"""
    ctx = _ctx(tmp_path)
    w = await FileWriteTool().execute(
        {"path": "out.md", "content": "# Title\n\n## Sec\nbody line\n"}, ctx
    )
    assert w.success is True
    assert ctx.landed_artifact_authors.get("out.md") == "a"
    ok = await FileReadTool().execute({"path": "out.md"}, ctx)
    assert ok.success is True
    assert "body line" in (ok.output or "")


async def test_file_read_author_and_reader_share_landed_ledger(tmp_path: Path):
    """同 execution 共享 landed 表；作者与读者都能读已落盘文件。"""
    author_ctx = _ctx(tmp_path, agent_id="writer")
    w = await FileWriteTool().execute(
        {"path": "shared.md", "content": "# Shared\n\nbody for downstream\n"},
        author_ctx,
    )
    assert w.success is True
    assert author_ctx.landed_artifact_kinds.get("shared.md") is not None

    reader_ctx = replace(author_ctx, agent_id="ceo", run_id="ceo-run")
    assert reader_ctx.landed_artifact_kinds is author_ctx.landed_artifact_kinds

    allowed = await FileReadTool().execute({"path": "shared.md"}, reader_ctx)
    assert allowed.success is True
    assert "body for downstream" in allowed.output
    still = await FileReadTool().execute({"path": "shared.md"}, author_ctx)
    assert still.success is True
    assert "body for downstream" in (still.output or "")


# --- file_write overwrite: no completeness heuristic ---


async def test_write_overwrite_omission_lands_without_nudge(tmp_path: Path):
    old = "A" * 200 + "\n完整中段内容\n" + "B" * 200
    assert len(old) >= 400
    (tmp_path / "draft.md").write_text(old, encoding="utf-8")
    truncated = (
        "A" * 40
        + "\n……（中间省略，已保留首尾）……\n"
        + "B" * 40
    )
    result = await FileWriteTool().execute(
        {"path": "draft.md", "content": truncated}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "draft.md").read_text(encoding="utf-8") == truncated
    assert "产物疑似不完整" not in (result.output or "")


async def test_write_overwrite_severe_shrink_lands(tmp_path: Path):
    old = "字" * 2000
    short = "字" * 300
    (tmp_path / "报告.md").write_text(old, encoding="utf-8")
    result = await FileWriteTool().execute(
        {"path": "报告.md", "content": short}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "报告.md").read_text(encoding="utf-8") == short
    assert "产物疑似不完整" not in (result.output or "")
    assert "allow_shrink" not in FileWriteTool().schema.parameters["properties"]


async def test_write_no_nudge_on_new_file(tmp_path: Path):
    # New file — even with omission-looking text — must not false-positive.
    body = "开头\n……（中间省略，已保留首尾）……\n结尾"
    result = await FileWriteTool().execute(
        {"path": "new.md", "content": body}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "产物疑似不完整" not in result.output


async def test_write_no_nudge_on_modest_edit(tmp_path: Path):
    old = "字" * 500
    assert len(old) >= 400
    (tmp_path / "essay.md").write_text(old, encoding="utf-8")
    # ~80% of old length, no omission markers — normal small revision.
    modest = "字" * 400
    result = await FileWriteTool().execute(
        {"path": "essay.md", "content": modest}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "产物疑似不完整" not in result.output


def test_has_omission_marker_covers_en_and_cn():
    from agentcore.tools.builtin.file_ops import has_omission_marker

    assert has_omission_marker("……（中间省略，已保留首尾）……")
    assert has_omission_marker("正文（略）续")
    assert has_omission_marker("see ... omitted details")
    assert has_omission_marker("Truncated for brevity here")
    assert not has_omission_marker("正常全文无省略")


# --- file_write / append / str_replace: no hard length gate ---


async def test_write_allows_oversized_prose(tmp_path: Path):
    """超阈值正文一次 file_write 须成功落盘（不再硬拒）。"""
    # Former hard-gate threshold was ≈2000 tokens × 4 chars ≈ 8000 chars.
    body = "x" * 8000
    result = await FileWriteTool().execute(
        {"path": "big.html", "content": body}, _ctx(tmp_path)
    )
    assert result.success is True
    assert result.contract_failure is not True
    assert (tmp_path / "big.html").read_text(encoding="utf-8") == body
    assert "拒绝整篇一次写入" not in (result.error or "")
    assert "拒绝整篇一次写入" not in (result.output or "")
    assert "已拦截" not in (result.error or "")


async def test_append_allows_oversized_chunk(tmp_path: Path):
    from agentcore.tools.builtin.file_ops import FileAppendTool

    (tmp_path / "a.md").write_text("# skeleton\n", encoding="utf-8")
    body = "y" * 8000
    result = await FileAppendTool().execute(
        {"path": "a.md", "content": body}, _ctx(tmp_path)
    )
    assert result.success is True
    assert result.contract_failure is not True
    assert "拒绝单次过大写入" not in (result.error or "")
    assert (tmp_path / "a.md").read_text(encoding="utf-8") == "# skeleton\n" + body


async def test_str_replace_allows_oversized_new_string(tmp_path: Path):
    from agentcore.tools.builtin.file_ops import StrReplaceTool

    (tmp_path / "a.md").write_text("## 参考文献\n", encoding="utf-8")
    new = "z" * 8000
    result = await StrReplaceTool().execute(
        {
            "path": "a.md",
            "old_string": "## 参考文献",
            "new_string": new,
        },
        _ctx(tmp_path),
    )
    assert result.success is True
    assert result.contract_failure is not True
    assert "拒绝单次过大写入" not in (result.error or "")
    assert (tmp_path / "a.md").read_text(encoding="utf-8") == new + "\n"


async def test_write_allows_medium_prose_body(tmp_path: Path):
    body = "x" * 4000
    result = await FileWriteTool().execute(
        {"path": "ok.md", "content": body}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "ok.md").read_text(encoding="utf-8") == body
    assert "拒绝整篇一次写入" not in (result.output or "")


async def test_write_allows_short_skeleton_with_section_markers(tmp_path: Path):
    skeleton = (
        "# Outline\n\n<!-- OUTLINE -->\n"
        "<!-- SECTION:s0 START -->\n<!-- SECTION:s0 END -->\n"
    )
    result = await FileWriteTool().execute(
        {"path": "report.md", "content": skeleton}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "kind: skeleton" in result.output


async def test_write_then_append_segmented_path(tmp_path: Path):
    """建站 HTML 短骨架 + SECTION 填空：append 仍放行（勿误伤）。"""
    skeleton = (
        "<!doctype html>\n<html>\n<head></head>\n<body>\n"
        "<!-- SECTION:s0 START -->\n<!-- SECTION:s0 END -->\n"
    )
    section = "  <section>hello</section>\n"
    closing = "</body>\n</html>\n"

    ctx = _ctx(tmp_path)
    w = await FileWriteTool().execute(
        {"path": "site/index.html", "content": skeleton}, ctx
    )
    assert w.success is True
    assert ctx.landed_artifact_kinds.get("site/index.html") == "skeleton"

    a1 = await FileAppendTool().execute(
        {"path": "site/index.html", "content": section}, ctx
    )
    assert a1.success is True
    assert "已追加" in a1.output

    a2 = await FileAppendTool().execute(
        {"path": "site/index.html", "content": closing}, ctx
    )
    assert a2.success is True
    merged = (tmp_path / "site" / "index.html").read_text(encoding="utf-8")
    assert merged == skeleton + section + closing


def test_write_schema_does_not_teach_completeness_gates():
    """按钮只留这是什么 + HOW；完整性硬拒不进 schema / landing。"""
    from agentcore.runtime.skills import build_system_skill_registry

    write_desc = FileWriteTool().schema.description
    assert "写入文件" in write_desc
    assert "HOW→consult(long_form_landing)" in write_desc
    assert "主路径" not in write_desc
    assert "省略标记" not in write_desc
    assert "50%" not in write_desc
    assert "800" not in write_desc
    assert "括号" not in write_desc
    assert "硬拒" not in write_desc
    assert "清参后改稿" not in write_desc
    landing = build_system_skill_registry().get("long_form_landing")
    assert landing is not None
    assert "省略标记" not in landing.body
    assert "allow_shrink" not in landing.body
    assert "硬拒" not in landing.body
    assert "HOW→consult(long_form_landing)" in FileAppendTool().schema.description
    assert "HOW→consult(long_form_landing)" in StrReplaceTool().schema.description
    content_desc = FileWriteTool().schema.parameters["properties"]["content"]["description"]
    assert "完整正文" in content_desc
    assert "硬拒" not in content_desc
    write_path = FileWriteTool().schema.parameters["properties"]["path"]["description"]
    append_path = FileAppendTool().schema.parameters["properties"]["path"]["description"]
    replace_path = StrReplaceTool().schema.parameters["properties"]["path"]["description"]
    assert "扁平" in write_path
    assert "扁平" not in append_path
    assert "扁平" not in replace_path

    append_desc = FileAppendTool().schema.description
    assert "追加" in append_desc
    assert "骨架" not in append_desc
    assert "次数上限" not in append_desc
    assert "Artifact-first" not in append_desc

    replace_desc = StrReplaceTool().schema.description
    assert "完全匹配" in replace_desc or "精确替换" in replace_desc
    assert "清参后改稿" not in replace_desc
    assert "_landed_summary" not in replace_desc
    assert "Artifact-first" not in replace_desc
    new_desc = StrReplaceTool().schema.parameters["properties"]["new_string"]["description"]
    assert "不硬拒" in new_desc
    assert "_landed_summary" not in new_desc
    assert "已落盘短状态" in new_desc or "清理占位" in new_desc


def test_classify_write_kind_helpers():
    from agentcore.tools.builtin.file_ops import (
        classify_write_kind,
        extract_title_tree,
        has_skeleton_markers,
    )

    assert has_skeleton_markers("<!-- SECTION:s0 START -->")
    assert has_skeleton_markers("<!-- OUTLINE -->")
    assert classify_write_kind("# A\n\n## B\n\n<!-- OUTLINE -->\n") == "skeleton"
    assert classify_write_kind("短") == "skeleton"
    prose = "# T\n\n" + ("正文内容。" * 80)
    assert classify_write_kind(prose) == "prose"
    assert extract_title_tree("# Hello\n\n## World\n") == ["# Hello", "## World"]


# --- file_delete ---


def test_delete_schema_is_short_trigger():
    """恢复路径在成功回执，不进每轮 schema。"""
    schema = FileDeleteTool().schema
    blob = schema.description + schema.parameters["properties"]["permanent"]["description"]
    assert "可逆" in schema.description
    assert "permanent" in schema.description
    assert "系统回收站" not in blob
    assert "AgentCore/trash" not in blob
    assert "一键还原" not in blob


async def test_delete_file(tmp_path: Path):
    (tmp_path / "f.txt").write_text("bye", encoding="utf-8")
    result = await FileDeleteTool().execute({"path": "f.txt"}, _ctx(tmp_path))
    assert result.success is True
    assert "可逆删除" in result.output
    assert "系统回收站" in result.output or "AgentCore/trash" in result.output
    assert not (tmp_path / "f.txt").exists()
    # Soft-deleted into workspace trash with restore metadata.
    trash = tmp_path / "AgentCore" / "trash"
    assert trash.is_dir()
    entries = list(trash.iterdir())
    assert len(entries) == 1
    assert (entries[0] / "meta.json").is_file()
    assert (entries[0] / "content").read_text(encoding="utf-8") == "bye"


async def test_delete_directory_recursive(tmp_path: Path):
    (tmp_path / "pkg").mkdir()
    (tmp_path / "pkg" / "a.txt").write_text("a", encoding="utf-8")
    (tmp_path / "pkg" / "sub").mkdir()
    (tmp_path / "pkg" / "sub" / "b.txt").write_text("b", encoding="utf-8")
    result = await FileDeleteTool().execute({"path": "pkg"}, _ctx(tmp_path))
    assert result.success is True
    assert not (tmp_path / "pkg").exists()
    assert (tmp_path / "AgentCore" / "trash").is_dir()


async def test_delete_permanent_hard_removes(tmp_path: Path):
    (tmp_path / "f.txt").write_text("bye", encoding="utf-8")
    result = await FileDeleteTool().execute(
        {"path": "f.txt", "permanent": True}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "永久删除" in result.output
    assert not (tmp_path / "f.txt").exists()
    trash = tmp_path / "AgentCore" / "trash"
    assert not trash.exists() or not any(trash.iterdir())


async def test_delete_allows_substantial_draft(tmp_path: Path):
    body = "成篇正文。" * 80
    (tmp_path / "report.md").write_text(body, encoding="utf-8")
    result = await FileDeleteTool().execute({"path": "report.md"}, _ctx(tmp_path))
    assert result.success is True
    assert not (tmp_path / "report.md").exists()
    assert "可逆删除" in result.output


async def test_delete_not_found(tmp_path: Path):
    result = await FileDeleteTool().execute({"path": "nope.txt"}, _ctx(tmp_path))
    assert result.success is False
    assert "路径不存在" in result.error
    assert result.contract_failure is True


async def test_delete_rejects_path_outside_workspace(tmp_path: Path):
    ws = tmp_path / "ws"
    ws.mkdir()
    (tmp_path / "secret.txt").write_text("top secret", encoding="utf-8")
    result = await FileDeleteTool().execute({"path": "../secret.txt"}, _ctx(ws))
    assert result.success is False
    assert "超出了工作区范围" in result.error
    # the out-of-tree file must be untouched
    assert (tmp_path / "secret.txt").read_text(encoding="utf-8") == "top secret"


async def test_delete_refuses_workspace_root(tmp_path: Path):
    (tmp_path / "keep.txt").write_text("keep", encoding="utf-8")
    result = await FileDeleteTool().execute({"path": ""}, _ctx(tmp_path))
    assert result.success is False
    # Empty path is rejected up-front; "." still hits OutsideWorkspace at the
    # backend for genuine root deletes.
    assert "path 不能为空" in result.error or "超出了工作区范围" in result.error
    # nothing in the root was removed
    assert (tmp_path / "keep.txt").exists()


async def test_delete_refuses_dot_workspace_root(tmp_path: Path):
    (tmp_path / "keep.txt").write_text("keep", encoding="utf-8")
    result = await FileDeleteTool().execute({"path": "."}, _ctx(tmp_path))
    assert result.success is False
    assert "超出了工作区范围" in result.error or "工作区根" in (result.error or "")
    assert (tmp_path / "keep.txt").exists()


# --- file_move ---


async def test_move_renames_file(tmp_path: Path):
    (tmp_path / "old.txt").write_text("data", encoding="utf-8")
    result = await FileMoveTool().execute(
        {"source": "old.txt", "destination": "new.txt"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "已把 old.txt 移动到 new.txt" in result.output
    assert not (tmp_path / "old.txt").exists()
    assert (tmp_path / "new.txt").read_text(encoding="utf-8") == "data"


async def test_move_creates_destination_parents(tmp_path: Path):
    (tmp_path / "f.txt").write_text("x", encoding="utf-8")
    result = await FileMoveTool().execute(
        {"source": "f.txt", "destination": "deep/nested/f.txt"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "deep" / "nested" / "f.txt").read_text(encoding="utf-8") == "x"


async def test_move_directory(tmp_path: Path):
    (tmp_path / "src").mkdir()
    (tmp_path / "src" / "a.txt").write_text("a", encoding="utf-8")
    result = await FileMoveTool().execute({"source": "src", "destination": "dst"}, _ctx(tmp_path))
    assert result.success is True
    assert (tmp_path / "dst" / "a.txt").read_text(encoding="utf-8") == "a"
    assert not (tmp_path / "src").exists()


async def test_move_refuses_to_overwrite(tmp_path: Path):
    (tmp_path / "a.txt").write_text("from", encoding="utf-8")
    (tmp_path / "b.txt").write_text("to", encoding="utf-8")
    result = await FileMoveTool().execute(
        {"source": "a.txt", "destination": "b.txt"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert "已存在" in result.error
    # both files must be untouched
    assert (tmp_path / "a.txt").read_text(encoding="utf-8") == "from"
    assert (tmp_path / "b.txt").read_text(encoding="utf-8") == "to"


async def test_move_source_not_found(tmp_path: Path):
    result = await FileMoveTool().execute(
        {"source": "ghost.txt", "destination": "x.txt"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert "源路径不存在" in result.error
    assert result.contract_failure is True


async def test_move_rejects_path_outside_workspace(tmp_path: Path):
    ws = tmp_path / "ws"
    ws.mkdir()
    (ws / "inside.txt").write_text("inside", encoding="utf-8")
    result = await FileMoveTool().execute(
        {"source": "inside.txt", "destination": "../escaped.txt"}, _ctx(ws)
    )
    assert result.success is False
    assert "超出了工作区范围" in result.error
    assert (ws / "inside.txt").read_text(encoding="utf-8") == "inside"
    assert not (tmp_path / "escaped.txt").exists()


async def test_move_requires_both_args(tmp_path: Path):
    (tmp_path / "f.txt").write_text("x", encoding="utf-8")
    result = await FileMoveTool().execute({"source": "f.txt"}, _ctx(tmp_path))
    assert result.success is False
    assert "必填" in result.error


async def test_move_identical_paths_is_idempotent(tmp_path: Path):
    (tmp_path / "f.txt").write_text("x", encoding="utf-8")
    result = await FileMoveTool().execute(
        {"source": "f.txt", "destination": "f.txt"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "相同" in result.output or "无需" in result.output
    assert (tmp_path / "f.txt").read_text(encoding="utf-8") == "x"


async def test_move_identical_after_dossier_flatten(tmp_path: Path):
    """Source already flat; nested reviews dest sanitizes to same path → idempotent."""
    flat = f"{REVIEWS_PREFIX}a_b_c.md"
    nested = f"{REVIEWS_PREFIX}a/b/c.md"
    flat_path = tmp_path.joinpath(*flat.split("/"))
    flat_path.parent.mkdir(parents=True, exist_ok=True)
    flat_path.write_text("review", encoding="utf-8")

    result = await FileMoveTool().execute(
        {"source": flat, "destination": nested}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "相同" in result.output or "无需" in result.output
    assert flat_path.read_text(encoding="utf-8") == "review"
    assert flat_path.exists()
    assert not tmp_path.joinpath(*nested.split("/")).exists()


# --- file_copy / mkdir / file_batch ---


async def test_copy_file_and_tree(tmp_path: Path):
    (tmp_path / "a.txt").write_text("data", encoding="utf-8")
    result = await FileCopyTool().execute(
        {"source": "a.txt", "destination": "b/c.txt"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "a.txt").read_text(encoding="utf-8") == "data"
    assert (tmp_path / "b" / "c.txt").read_text(encoding="utf-8") == "data"

    (tmp_path / "tree" / "sub").mkdir(parents=True)
    (tmp_path / "tree" / "sub" / "x.bin").write_bytes(b"\x00\xff")
    result = await FileCopyTool().execute(
        {"source": "tree", "destination": "tree2"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert (tmp_path / "tree2" / "sub" / "x.bin").read_bytes() == b"\x00\xff"


async def test_copy_identical_paths_is_idempotent(tmp_path: Path):
    (tmp_path / "f.txt").write_text("x", encoding="utf-8")
    result = await FileCopyTool().execute(
        {"source": "f.txt", "destination": "f.txt"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "相同" in result.output or "无需" in result.output
    assert (tmp_path / "f.txt").read_text(encoding="utf-8") == "x"


async def test_copy_identical_after_dossier_flatten(tmp_path: Path):
    """Source already flat; nested reviews dest sanitizes to same path → idempotent."""
    flat = f"{REVIEWS_PREFIX}a_b_c.md"
    nested = f"{REVIEWS_PREFIX}a/b/c.md"
    flat_path = tmp_path.joinpath(*flat.split("/"))
    flat_path.parent.mkdir(parents=True, exist_ok=True)
    flat_path.write_text("review", encoding="utf-8")

    result = await FileCopyTool().execute(
        {"source": flat, "destination": nested}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "相同" in result.output or "无需" in result.output
    assert flat_path.read_text(encoding="utf-8") == "review"
    assert flat_path.exists()
    assert not tmp_path.joinpath(*nested.split("/")).exists()


async def test_copy_refuses_overwrite(tmp_path: Path):
    (tmp_path / "a.txt").write_text("from", encoding="utf-8")
    (tmp_path / "b.txt").write_text("to", encoding="utf-8")
    result = await FileCopyTool().execute(
        {"source": "a.txt", "destination": "b.txt"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert "已存在" in result.error


def test_mkdir_schema_teaches_structure_not_app_shell():
    desc = MkdirTool().schema.description
    assert "结构目录" in desc
    assert "src/" in desc
    assert "不必先 mkdir" in desc
    assert "套应用名/话题名当工程根" in desc
    assert "≠" in desc
    assert "whiteboard" not in desc
    assert "court-game" not in desc
    assert "禁止" not in desc


async def test_mkdir_creates_and_refuses_existing(tmp_path: Path):
    result = await MkdirTool().execute({"path": "out/docs"}, _ctx(tmp_path))
    assert result.success is True
    assert (tmp_path / "out" / "docs").is_dir()
    result = await MkdirTool().execute({"path": "out/docs"}, _ctx(tmp_path))
    assert result.success is False
    assert "已存在" in result.error


async def test_file_batch_partial_failure_continues(tmp_path: Path):
    (tmp_path / "a.txt").write_text("a", encoding="utf-8")
    result = await FileBatchTool().execute(
        {
            "operations": [
                {"op": "mkdir", "path": "out"},
                {"op": "copy", "source": "a.txt", "destination": "out/a.txt"},
                {"op": "move", "source": "missing.txt", "destination": "out/m.txt"},
                {"op": "delete", "path": "ghost.txt"},
                {"op": "mkdir", "path": "out"},  # already exists → skip
            ]
        },
        _ctx(tmp_path),
    )
    assert result.success is False  # one hard failure (move missing)
    assert "本次共 5 项" in result.output
    assert "成功" in result.output
    assert "跳过" in result.output
    assert "失败" in result.output
    assert (tmp_path / "out" / "a.txt").read_text(encoding="utf-8") == "a"
    assert result.metadata["ok"] >= 2
    assert result.metadata["fail"] >= 1


# --- file_list / glob ---


def test_expand_brace_globs_basic():
    assert expand_brace_globs("*.{ts,tsx}") == ["*.ts", "*.tsx"]
    assert expand_brace_globs("**/*.{py,pyi}") == ["**/*.py", "**/*.pyi"]
    assert expand_brace_globs("*.py") == ["*.py"]
    assert expand_brace_globs("*") == ["*"]


def test_compile_glob_pattern_globstar():
    assert compile_glob_pattern("") is None
    assert compile_glob_pattern("*.py") == GlobPlan(None, ".", "*.py", GLOB_DEPTH)
    assert compile_glob_pattern("**/*.py") == GlobPlan(None, ".", "*.py", GLOB_DEPTH)
    assert compile_glob_pattern("*") == GlobPlan(None, ".", "*", GLOB_DEPTH)
    assert compile_glob_pattern("**/*") == GlobPlan(None, ".", "*", GLOB_DEPTH)
    assert compile_glob_pattern("src/*.py") == GlobPlan(None, "src", "*.py", 1)
    assert compile_glob_pattern("src/**/*.py") == GlobPlan(
        None, "src", "*.py", GLOB_DEPTH
    )
    assert compile_glob_pattern("**/observability/**") == GlobPlan(
        "observability", ".", "*", GLOB_DEPTH
    )
    assert compile_glob_pattern("packages/*/package.json") == GlobPlan(
        None, "packages", "package.json", 1, True
    )
    assert compile_glob_pattern("packages/*/*.package.json") == GlobPlan(
        None, "packages", "*.package.json", 1, True
    )
    assert compile_glob_pattern("*/package.json") == GlobPlan(
        None, ".", "package.json", 1, True
    )
    assert compile_glob_pattern("packages/*/*") == GlobPlan(
        None, "packages", "*", 1, True
    )
    assert compile_glob_patterns("*.{ts,tsx}") == [
        GlobPlan(None, ".", "*.ts", GLOB_DEPTH),
        GlobPlan(None, ".", "*.tsx", GLOB_DEPTH),
    ]


def test_file_list_schema_is_one_layer_ls():
    schema = FileListTool().schema
    props = schema.parameters["properties"]
    assert "directory" in props
    assert "pattern" not in props
    assert "recursive" not in props
    assert "max_depth" not in props
    assert "glob" in schema.description


def test_glob_schema_requires_pattern():
    schema = GlobTool().schema
    assert schema.parameters["required"] == ["pattern"]
    props = schema.parameters["properties"]
    assert "pattern" in props
    assert "path" in props
    assert "directory" not in props
    assert "recursive" not in props
    assert "pkg/*/name" in props["pattern"]["description"]
    assert "pkg/*/name" in schema.description


async def test_glob_finds_nested_files_from_root(tmp_path: Path):
    (tmp_path / "server").mkdir()
    (tmp_path / "server" / "main.py").write_text("x", encoding="utf-8")
    (tmp_path / "client").mkdir()
    (tmp_path / "package.json").write_text("{}", encoding="utf-8")

    result = await GlobTool().execute({"pattern": "*.py"}, _ctx(tmp_path))
    assert result.success is True
    assert "空目录" not in (result.output or "")
    assert "server/main.py" in (result.output or "")
    assert "├──" not in (result.output or "")


@pytest.mark.parametrize(
    "extra",
    [{"pattern": "*.py"}, {"recursive": True}, {"max_depth": 4}],
)
async def test_file_list_leftover_fails_to_glob(tmp_path: Path, extra: dict):
    result = await FileListTool().execute({"directory": ".", **extra}, _ctx(tmp_path))
    assert result.success is False
    assert result.contract_failure is True
    assert "glob" in (result.error or "")
    for key in extra:
        assert key in (result.error or "")


async def test_glob_star_lists_nested_files(tmp_path: Path):
    nested = tmp_path / "apps" / "server"
    nested.mkdir(parents=True)
    (nested / "main.py").write_text("x", encoding="utf-8")
    (tmp_path / "README.md").write_text("desk\n", encoding="utf-8")

    result = await GlobTool().execute({"pattern": "**/*"}, _ctx(tmp_path))
    assert result.success is True
    assert "apps/server/main.py" in (result.output or "")
    assert "README.md" in (result.output or "")


async def test_glob_src_star_py_is_one_level(tmp_path: Path):
    src = tmp_path / "src"
    src.mkdir()
    (src / "a.py").write_text("x", encoding="utf-8")
    (src / "nested").mkdir()
    (src / "nested" / "b.py").write_text("x", encoding="utf-8")

    result = await GlobTool().execute({"pattern": "src/*.py"}, _ctx(tmp_path))
    assert result.success is True
    assert "src/a.py" in (result.output or "")
    assert "nested/b.py" not in (result.output or "")


async def test_glob_star_dir_segment_matches_one_level_children(tmp_path: Path):
    """``packages/*/package.json`` expands child dirs; does not treat ``*`` as a name."""
    core = tmp_path / "packages" / "core"
    app = tmp_path / "packages" / "app"
    core.mkdir(parents=True)
    app.mkdir(parents=True)
    (core / "package.json").write_text("{}", encoding="utf-8")
    (app / "package.json").write_text("{}", encoding="utf-8")
    (tmp_path / "packages" / "package.json").write_text("{}", encoding="utf-8")
    nested = core / "nested"
    nested.mkdir()
    (nested / "package.json").write_text("{}", encoding="utf-8")

    result = await GlobTool().execute(
        {"pattern": "packages/*/package.json"}, _ctx(tmp_path)
    )
    assert result.success is True
    out = result.output or ""
    lines = [
        line.strip().replace("\\", "/")
        for line in out.splitlines()
        if line.strip() and not line.startswith("（")
    ]
    assert "packages/core/package.json" in lines
    assert "packages/app/package.json" in lines
    assert "packages/package.json" not in lines
    assert not any(line.endswith("nested/package.json") for line in lines)


async def test_glob_listing_error_names_failed_path_not_search_root(tmp_path: Path):
    """list_tree 失败应报后端拒绝的路径，而不是 glob 搜索根 ``.``。"""
    from agentcore.workspace.protocol import NotADirectory

    ctx = _ctx(tmp_path)

    async def boom(directory: str, **_kwargs: object):
        raise NotADirectory("packages/*")

    ctx.backend.list_tree = boom  # type: ignore[method-assign]
    result = await GlobTool().execute({"pattern": "*.py"}, ctx)
    assert result.success is False
    assert result.error is not None
    assert result.error.startswith("不是目录：packages/*")
    assert not result.error.startswith("不是目录：.")


async def test_glob_src_recursive_py(tmp_path: Path):
    src = tmp_path / "src"
    src.mkdir()
    (src / "a.py").write_text("x", encoding="utf-8")
    (src / "nested").mkdir()
    (src / "nested" / "b.py").write_text("x", encoding="utf-8")

    result = await GlobTool().execute({"pattern": "src/**/*.py"}, _ctx(tmp_path))
    assert result.success is True
    assert "src/a.py" in (result.output or "")
    assert "src/nested/b.py" in (result.output or "")


async def test_glob_any_named_dir_contents(tmp_path: Path):
    d = tmp_path / "apps" / "server" / "observability"
    d.mkdir(parents=True)
    (d / "catalog.py").write_text("x", encoding="utf-8")

    result = await GlobTool().execute(
        {"pattern": "**/observability/**"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "catalog.py" in (result.output or "")


async def test_glob_missing_path_finds_named_dir(tmp_path: Path):
    d = tmp_path / "apps" / "server" / "observability"
    d.mkdir(parents=True)
    (d / "a.md").write_text("x", encoding="utf-8")

    result = await GlobTool().execute(
        {"path": "observability", "pattern": "*.md"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "a.md" in (result.output or "")
    assert "不存在" in (result.output or "")


async def test_glob_directory_alias(tmp_path: Path):
    src = tmp_path / "src"
    src.mkdir()
    (src / "a.py").write_text("x", encoding="utf-8")
    result = await GlobTool().execute(
        {"pattern": "*.py", "directory": "src"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "src/a.py" in (result.output or "")


async def test_glob_empty_pattern_rejected(tmp_path: Path):
    result = await GlobTool().execute({"pattern": "  "}, _ctx(tmp_path))
    assert result.success is False
    assert result.contract_failure is True


async def test_glob_leftover_recursive_rejected(tmp_path: Path):
    result = await GlobTool().execute(
        {"pattern": "*.py", "recursive": True}, _ctx(tmp_path)
    )
    assert result.success is False
    assert result.contract_failure is True
    assert "recursive" in (result.error or "")


async def test_file_list_star_stays_one_layer(tmp_path: Path):
    nested = tmp_path / "apps" / "server" / "agentcore"
    nested.mkdir(parents=True)
    (nested / "main.py").write_text("x", encoding="utf-8")
    (tmp_path / "README.md").write_text("desk\n", encoding="utf-8")

    result = await FileListTool().execute({"directory": "."}, _ctx(tmp_path))
    assert result.success is True
    assert "apps" in result.output
    assert "main.py" not in result.output


async def test_file_list_truly_empty_dir_still_says_empty(tmp_path: Path):
    empty = tmp_path / "blank"
    empty.mkdir()
    result = await FileListTool().execute({"directory": "blank"}, _ctx(tmp_path))
    assert result.success is True
    assert "空目录" in result.output
    assert "无匹配" not in result.output


async def test_glob_brace_matches_either_extension(tmp_path: Path):
    src = tmp_path / "client" / "src"
    src.mkdir(parents=True)
    (src / "App.tsx").write_text("export {}", encoding="utf-8")
    (src / "api.ts").write_text("export {}", encoding="utf-8")
    (src / "readme.md").write_text("x", encoding="utf-8")

    result = await GlobTool().execute(
        {"path": "client/src", "pattern": "*.{ts,tsx}"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "空目录" not in (result.output or "")
    assert "App.tsx" in (result.output or "")
    assert "api.ts" in (result.output or "")
    assert "readme.md" not in (result.output or "")


async def test_glob_miss_hint_does_not_claim_empty(tmp_path: Path):
    (tmp_path / "pkg").mkdir()
    (tmp_path / "pkg" / "a.py").write_text("x", encoding="utf-8")
    result = await GlobTool().execute(
        {"path": "pkg", "pattern": "*.rs"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "空目录" not in (result.output or "")
    assert "无匹配 pattern='*.rs'" in (result.output or "")
    assert "a.py" in (result.output or "") or "目录非空" in (result.output or "")


async def test_file_read_missing_with_parent_gives_landmark(tmp_path: Path):
    """父目录存在时：同层样本 + 更宽查找建议（按名走 glob）。"""
    app = tmp_path / "apps" / "desktop"
    app.mkdir(parents=True)
    (app / "README.md").write_text("# desk", encoding="utf-8")
    (app / "src").mkdir()

    result = await FileReadTool().execute(
        {"path": "apps/desktop/package.json"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert result.error is not None
    assert result.error.startswith("文件不存在：apps/desktop/package.json")
    assert result.failure_code == "not_found"
    assert result.failure_message is not None
    assert result.failure_message.startswith("文件不存在：apps/desktop/package.json")
    assert result.contract_failure is True
    assert "父目录" in result.error
    assert "apps/desktop/" in result.error
    assert "可见同层示例" in result.error
    assert "README.md" in result.error or "src" in result.error
    assert "glob" in result.error
    assert "file_list(pattern)" not in result.error
    assert "更宽查找" in result.error
    assert "已知路径" in result.error
    assert "反复重试" in result.error


async def test_file_list_missing_with_parent_gives_landmark(tmp_path: Path):
    """列目录路径不存在但上级可列：同层样本 + 勿反复重试。"""
    server = tmp_path / "apps" / "server"
    server.mkdir(parents=True)
    (server / "agentcore").mkdir()
    (server / "README.md").write_text("x", encoding="utf-8")
    result = await FileListTool().execute(
        {"directory": "apps/server/src"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert result.error is not None
    assert result.error.startswith("不是目录：apps/server/src")
    assert result.failure_code == "not_found"
    assert result.contract_failure is True
    assert "父目录" in result.error
    assert "apps/server/" in result.error
    assert "agentcore" in result.error or "README.md" in result.error
    assert "反复重试" in result.error


async def test_file_read_missing_parent_gives_root_tip(tmp_path: Path):
    """路径与父目录都不在：不编造同层样本，但给根查找 / 勿反复重试提示。"""
    (tmp_path / "apps").mkdir()
    result = await FileReadTool().execute(
        {"path": "apps/ghost/package.json"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert result.error is not None
    assert result.error.startswith("文件不存在：apps/ghost/package.json")
    assert result.contract_failure is True
    assert "上级目录也找不到" in result.error
    assert "禁止凭通用目录名" not in result.error
    assert "反复重试" in result.error
    assert "可见同层示例" not in result.error
    assert "glob" in result.error
    assert "file_list(pattern)" not in result.error


async def test_file_list_latent_stage_dir_returns_empty_not_error(tmp_path: Path):
    """约定出口尚未创建：file_list 成功空态（写入会自动创建），不报 NotADirectory。"""
    from agentcore.workspace.stage_dirs import RESEARCH_DIR

    assert not (tmp_path / "AgentCore").exists()
    result = await FileListTool().execute(
        {"directory": RESEARCH_DIR}, _ctx(tmp_path)
    )
    assert result.success is True
    assert result.error is None
    assert "空目录" in (result.output or "")
    assert "写入时会自动创建" in (result.output or "")
    assert not (tmp_path / "AgentCore").exists()  # 不预创建


async def test_file_list_latent_attachments_returns_empty_not_error(tmp_path: Path):
    """attachments/ 尚未创建：列目录空态成功。"""
    result = await FileListTool().execute(
        {"directory": "attachments"}, _ctx(tmp_path)
    )
    assert result.success is True
    assert "写入时会自动创建" in (result.output or "")
    assert not (tmp_path / "attachments").exists()


async def test_file_list_guessed_missing_path_still_errors(tmp_path: Path):
    """真·乱猜路径仍报错（不得因 latent 口径放成空目录）。"""
    result = await FileListTool().execute(
        {"directory": "apps/server/src"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert result.error is not None
    assert result.error.startswith("不是目录：apps/server/src")
    assert result.failure_code == "not_found"
    assert result.contract_failure is True
    assert "写入时会自动创建" not in result.error


async def test_file_read_missing_top_level_uses_root_landmark(tmp_path: Path):
    """顶层缺失文件：父目录为根，仍给同层样本。"""
    (tmp_path / "README.md").write_text("hi", encoding="utf-8")
    (tmp_path / "src").mkdir()
    result = await FileReadTool().execute({"path": "package.json"}, _ctx(tmp_path))
    assert result.success is False
    assert result.error is not None
    assert result.error.startswith("文件不存在：package.json")
    assert result.contract_failure is True
    assert "父目录 ./" in result.error
    assert "可见同层示例" in result.error
    assert "README.md" in result.error or "src" in result.error
    assert "glob" in result.error
    assert "file_list(pattern)" not in result.error


async def test_file_list_shows_attachment_zip_hides_elsewhere(tmp_path: Path):
    """attachments/ only-zip must not render「（空目录）」; root zip stays hidden."""
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "pack.zip").write_bytes(b"PK\x03\x04")
    (tmp_path / "noise.zip").write_bytes(b"PK\x03\x04")

    att = await FileListTool().execute(
        {"directory": "attachments"}, _ctx(tmp_path)
    )
    assert att.success is True
    assert "空目录" not in att.output
    assert "pack.zip" in att.output

    root = await FileListTool().execute({"directory": "."}, _ctx(tmp_path))
    assert root.success is True
    assert "noise.zip" not in root.output
    assert "attachments" in root.output

    found = await GlobTool().execute({"pattern": "*.zip"}, _ctx(tmp_path))
    assert found.success is True
    assert "attachments/pack.zip" in (found.output or "")
    assert "noise.zip" in (found.output or "")



async def test_file_list_reveals_material_png(tmp_path: Path):
    """Materials path (e.g. src/shot.png) visible; sibling AI-noise still hidden."""
    (tmp_path / "src").mkdir()
    (tmp_path / "src" / "shot.png").write_bytes(b"png")
    (tmp_path / "src" / "other.png").write_bytes(b"png")
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "pack.zip").write_bytes(b"PK\x03\x04")

    ctx = _ctx(tmp_path)
    ctx.material_paths = frozenset({"src/shot.png"})
    ctx.backend.ai_list_materials = ctx.material_paths

    listed = await FileListTool().execute({"directory": "src"}, ctx)
    assert listed.success is True
    assert "shot.png" in listed.output
    assert "other.png" not in listed.output

    found = await GlobTool().execute({"pattern": "*.png"}, ctx)
    assert found.success is True
    assert "src/shot.png" in (found.output or "")
    assert "other.png" not in (found.output or "")

    att = await FileListTool().execute({"directory": "attachments"}, ctx)
    assert att.success is True
    assert "pack.zip" in (att.output or "")

    # Without materials, same png stays hidden
    bare = await FileListTool().execute({"directory": "src"}, _ctx(tmp_path))
    assert bare.success is True
    assert "shot.png" not in bare.output
    assert "空目录" in bare.output or bare.output.strip() == "" or "shot" not in bare.output


async def test_file_list_external_mount_shows_archive_zip(tmp_path: Path):
    """external/<alias>/ 下列出可见用户压缩包；工作区根 zip 仍隐藏。"""
    from agentcore.workspace.external_mounts import ExternalMount

    ext = tmp_path / "ext_root"
    ext.mkdir()
    (ext / "咨询.sy.zip").write_bytes(b"PK\x03\x04")
    (ext / "note.txt").write_text("hi", encoding="utf-8")
    (tmp_path / "noise.zip").write_bytes(b"PK\x03\x04")

    ctx = _ctx(tmp_path)
    ctx.backend.attach_external_mounts(
        {
            "desk": ExternalMount(
                alias="desk",
                root_id="r",
                label="桌面",
                abs_path=str(ext),
                mode="readonly",
            )
        }
    )

    listed = await FileListTool().execute(
        {"directory": "external/desk"}, ctx
    )
    assert listed.success is True
    assert "咨询.sy.zip" in listed.output
    assert "note.txt" in listed.output

    root = await FileListTool().execute({"directory": "."}, ctx)
    assert root.success is True
    assert "noise.zip" not in root.output


async def test_file_list_pattern_zip_reveals_workspace_archive(tmp_path: Path):
    """pattern 指向压缩包后缀时，工作区根 zip 也应列出。"""
    (tmp_path / "pack.zip").write_bytes(b"PK\x03\x04")
    (tmp_path / "readme.md").write_text("x", encoding="utf-8")

    hidden = await FileListTool().execute({"directory": "."}, _ctx(tmp_path))
    assert hidden.success is True
    assert "pack.zip" not in hidden.output

    revealed = await GlobTool().execute({"pattern": "*.zip"}, _ctx(tmp_path))
    assert revealed.success is True
    assert "pack.zip" in (revealed.output or "")


async def test_file_list_bare_external_actionable_hint_no_mounts(tmp_path: Path):
    result = await FileListTool().execute(
        {"directory": "external"}, _ctx(tmp_path)
    )
    assert result.success is False
    assert "external/<别名>/" in result.error
    assert "尚无" in result.error or "授权" in result.error


async def test_file_list_bare_external_lists_current_mounts(tmp_path: Path):
    from agentcore.workspace.external_mounts import ExternalMount

    ctx = _ctx(tmp_path)
    ctx.backend.attach_external_mounts(
        {
            "desk": ExternalMount(
                alias="desk",
                root_id="r",
                label="桌面",
                abs_path=str(tmp_path),
                mode="readonly",
            )
        }
    )
    result = await FileListTool().execute({"directory": "external"}, ctx)
    assert result.success is False
    assert "external/<别名>/" in result.error
    assert "`external/desk/`" in result.error


# --- write_scope (冷启动 explore_memory) ---


def _explore_ctx(workspace: Path) -> ToolContext:
    ctx = _ctx(workspace)
    ctx.write_scope = "explore_memory"
    return ctx


async def test_write_scope_explore_memory_allows_research_note(tmp_path: Path):
    ctx = _explore_ctx(tmp_path)
    result = await FileWriteTool().execute(
        {
            "path": "AgentCore/文档/research/摸底笔记.md",
            "content": "# 笔记\n",
        },
        ctx,
    )
    assert result.success is True
    assert (tmp_path / "AgentCore" / "文档" / "research" / "摸底笔记.md").is_file()


async def test_write_scope_explore_memory_rejects_user_project_path(tmp_path: Path):
    ctx = _explore_ctx(tmp_path)
    result = await FileWriteTool().execute(
        {"path": "src/main.py", "content": "print(1)\n"},
        ctx,
    )
    assert result.success is False
    assert result.contract_failure is True
    assert "AgentCore/" in (result.error or "")
    assert "src/main.py" in (result.error or "")


async def test_write_scope_explore_memory_has_no_inner_path_ban(tmp_path: Path):
    """步 3：闸只判「在不在 AgentCore/ 下」——厚约定文档已是条目，worker 无工具可写。"""
    ctx = _explore_ctx(tmp_path)
    result = await FileWriteTool().execute(
        {
            "path": "AgentCore/文档/背景/架构详解.md",
            "content": "# 背景资料\n",
        },
        ctx,
    )
    assert result.success is True
    assert (tmp_path / "AgentCore" / "文档" / "背景" / "架构详解.md").is_file()


async def test_write_scope_none_rejects_all(tmp_path: Path):
    ctx = _ctx(tmp_path)
    ctx.write_scope = "none"
    result = await FileAppendTool().execute(
        {"path": "AgentCore/文档/research/x.md", "content": "x"},
        ctx,
    )
    assert result.success is False
    assert "write_scope=none" in (result.error or "")


async def test_write_scope_explore_memory_str_replace_rejects_outside(tmp_path: Path):
    (tmp_path / "app.py").write_text("a = 1\n", encoding="utf-8")
    ctx = _explore_ctx(tmp_path)
    result = await StrReplaceTool().execute(
        {"path": "app.py", "old_string": "a = 1", "new_string": "a = 2"},
        ctx,
    )
    assert result.success is False
    assert "AgentCore/" in (result.error or "")
