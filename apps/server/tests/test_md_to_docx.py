"""Tests for deterministic Markdown → DOCX conversion."""

from __future__ import annotations

import io
import zipfile
from pathlib import Path

import pytest
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

import agentcore.docs_export.md_to_docx as md_to_docx_mod
from agentcore.docs_export.md_to_docx import (
    collect_image_srcs,
    convert_markdown_to_docx,
    docx_path_for_markdown,
    resolve_workspace_image_path,
)
from agentcore.docs_export.workspace_export import ExportMarkdownError, export_markdown_path
from agentcore.tools.sandbox import SubprocessSandbox
from agentcore.workspace.server import ServerWorkspace


def test_python_docx_is_lazy_until_convert(monkeypatch: pytest.MonkeyPatch) -> None:
    """Importing the module must not require python-docx until first convert.

    Sidecar mis-bundles historically crashed chat at tool registration via eager
    ``from docx import …``. Capability still ships in both runtimes; lazy load is
    import hygiene so a missing stack fails at convert, not at pipeline import.
    """
    monkeypatch.setattr(md_to_docx_mod, "_DocumentFactory", None)
    # Re-bind other sentinels so _ensure_docx runs a real import.
    for name in (
        "_WD_ALIGN_PARAGRAPH",
        "_RT",
        "_qn",
        "_OxmlElement",
        "_Cm",
        "_Inches",
        "_Pt",
        "_RGBColor",
        "_MAX_IMAGE_WIDTH",
    ):
        monkeypatch.setattr(md_to_docx_mod, name, None)
    assert md_to_docx_mod._DocumentFactory is None
    out = convert_markdown_to_docx("# hi")
    assert md_to_docx_mod._DocumentFactory is not None
    assert out.docx_bytes[:2] == b"PK"



def _tiny_png() -> bytes:
    from PIL import Image

    buf = io.BytesIO()
    Image.new("RGB", (8, 8), color=(20, 120, 200)).save(buf, format="PNG")
    return buf.getvalue()


_FIXTURE_MD = """# 标题一

## 小节二

这是一段含 [链接](https://example.com) 与 **加粗** 的正文。

- 无序甲
- 无序乙

1. 有序一
2. 有序二

```python
print("hello")
```

| 列A | 列B |
| --- | --- |
| 1 | 2 |

![示意图](./assets/chart.png)

![缺失](./assets/missing.png)
"""


def test_docx_path_for_markdown():
    assert docx_path_for_markdown("报告.md") == "报告.docx"
    assert docx_path_for_markdown("docs/a.markdown") == "docs/a.docx"
    assert docx_path_for_markdown("noext") == "noext.docx"


def test_collect_and_resolve_images():
    srcs = collect_image_srcs(_FIXTURE_MD)
    assert "./assets/chart.png" in srcs
    assert "./assets/missing.png" in srcs
    assert resolve_workspace_image_path("docs/报告.md", "./assets/chart.png") == (
        "docs/assets/chart.png"
    )
    assert resolve_workspace_image_path("报告.md", "https://x/y.png") is None


def test_convert_markdown_structure_and_missing_image_warning():
    png = _tiny_png()
    result = convert_markdown_to_docx(
        _FIXTURE_MD,
        images={
            "./assets/chart.png": png,
            "./assets/missing.png": None,
        },
    )
    assert result.docx_bytes[:2] == b"PK"
    assert any("缺图" in w and "missing.png" in w for w in result.warnings)

    # OOXML package is a zip; document.xml must exist.
    with zipfile.ZipFile(io.BytesIO(result.docx_bytes)) as zf:
        names = zf.namelist()
        assert "word/document.xml" in names
        xml = zf.read("word/document.xml").decode("utf-8")
        assert "标题一" in xml
        assert "hello" in xml
        assert "列A" in xml

    doc = Document(io.BytesIO(result.docx_bytes))
    texts = [p.text for p in doc.paragraphs if p.text.strip()]
    assert any("标题一" in t for t in texts)
    assert any("无序甲" in t for t in texts)
    assert any("有序一" in t for t in texts)
    assert len(doc.tables) >= 1
    assert doc.tables[0].cell(0, 0).text.strip() == "列A"


_PLEADING_MD = """# 民事起诉状

原告：张三，男，1980 年生，住某市某区。

被告：李四，男，1979 年生，住某市某区。

## 诉讼请求

判令被告返还借款本金人民币十万元。

- 证据一
"""


def _document_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        return zf.read("word/document.xml").decode("utf-8")


def _footer_xml(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
        parts = [n for n in zf.namelist() if n.startswith("word/footer")]
        return "\n".join(zf.read(n).decode("utf-8") for n in parts)


def test_default_layout_centers_h1_without_first_line_indent():
    """默认档（技术文档口径）：A4、正文 1.5 倍行距、大标题居中，正文不缩进、不对齐、无页码。"""
    result = convert_markdown_to_docx(_PLEADING_MD)
    xml = _document_xml(result.docx_bytes)
    assert xml.count('<w:jc w:val="center"/>') == 1  # 仅 H1，H2 不居中
    assert "w:firstLine" not in xml
    assert '<w:jc w:val="both"/>' not in xml
    assert "PAGE" not in _footer_xml(result.docx_bytes)

    doc = Document(io.BytesIO(result.docx_bytes))
    section = doc.sections[0]
    assert section.page_width.cm == pytest.approx(21.0, abs=0.05)
    assert section.page_height.cm == pytest.approx(29.7, abs=0.05)
    assert section.top_margin.cm == pytest.approx(2.54, abs=0.05)
    assert section.left_margin.cm == pytest.approx(3.17, abs=0.05)
    h1 = next(p for p in doc.paragraphs if p.text.strip() == "民事起诉状")
    assert h1.alignment == WD_ALIGN_PARAGRAPH.CENTER
    body = next(p for p in doc.paragraphs if p.text.startswith("原告："))
    assert body.paragraph_format.first_line_indent is None
    assert body.alignment is None
    assert body.paragraph_format.line_spacing == pytest.approx(1.5)
    assert body.paragraph_format.space_after == Pt(3)
    item = next(p for p in doc.paragraphs if "证据一" in p.text)
    assert item.paragraph_format.line_spacing == pytest.approx(1.5)
    assert item.alignment is None or item.alignment == WD_ALIGN_PARAGRAPH.LEFT


def test_official_layout_indents_body_first_line():
    """公文档：正文首行缩进两字 + 两端对齐 + 公文页边距 + 页码；标题/列表不吃缩进与两端对齐。"""
    result = convert_markdown_to_docx(_PLEADING_MD, layout="official")
    xml = _document_xml(result.docx_bytes)
    assert xml.count('<w:jc w:val="center"/>') == 1
    # 三段正文（原告 / 被告 / 诉讼请求正文）各一条，标题与列表项不算。
    assert xml.count('<w:ind w:firstLine="480" w:firstLineChars="200"/>') == 3
    assert xml.count("w:firstLine=") == 3
    assert xml.count('<w:jc w:val="both"/>') == 3
    footer = _footer_xml(result.docx_bytes)
    assert "PAGE" in footer
    assert '<w:jc w:val="center"/>' in footer

    doc = Document(io.BytesIO(result.docx_bytes))
    section = doc.sections[0]
    assert section.page_width.cm == pytest.approx(21.0, abs=0.05)
    assert section.page_height.cm == pytest.approx(29.7, abs=0.05)
    assert section.top_margin.cm == pytest.approx(3.7, abs=0.05)
    assert section.bottom_margin.cm == pytest.approx(3.5, abs=0.05)
    assert section.left_margin.cm == pytest.approx(2.8, abs=0.05)
    assert section.right_margin.cm == pytest.approx(2.6, abs=0.05)
    body = next(p for p in doc.paragraphs if p.text.startswith("原告："))
    assert body.paragraph_format.first_line_indent == Pt(24)
    assert body.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY
    assert body.paragraph_format.line_spacing == pytest.approx(1.5)
    h1 = next(p for p in doc.paragraphs if p.text.strip() == "民事起诉状")
    assert h1.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert h1.paragraph_format.first_line_indent is None
    item = next(p for p in doc.paragraphs if "证据一" in p.text)
    assert item.alignment is None or item.alignment == WD_ALIGN_PARAGRAPH.LEFT
    assert item.paragraph_format.first_line_indent in (None, Pt(0))


def test_layout_never_inferred_from_body_text():
    """同一份公文正文，不传档位就必须与技术文档一个排版——禁止内容启发式。"""
    plain = _document_xml(convert_markdown_to_docx(_PLEADING_MD).docx_bytes)
    explicit = _document_xml(
        convert_markdown_to_docx(_PLEADING_MD, layout="standard").docx_bytes
    )
    assert plain == explicit
    assert "w:firstLine" not in plain


@pytest.mark.asyncio
async def test_export_markdown_path_layout_defaults_to_standard(tmp_path: Path):
    """HTTP「导出 Word」路径不传 layout，落盘文件必须没有首行缩进。"""
    root = tmp_path / "ws"
    root.mkdir()
    (root / "起诉状.md").write_text(_PLEADING_MD, encoding="utf-8")
    backend = ServerWorkspace(root=root, sandbox=SubprocessSandbox())

    await export_markdown_path(backend, "起诉状.md")
    assert "w:firstLine" not in _document_xml((root / "起诉状.docx").read_bytes())

    await export_markdown_path(backend, "起诉状.md", layout="official")
    assert "w:firstLine" in _document_xml((root / "起诉状.docx").read_bytes())


@pytest.mark.asyncio
async def test_export_markdown_path_writes_sibling(tmp_path: Path):
    root = tmp_path / "ws"
    root.mkdir()
    (root / "assets").mkdir()
    (root / "报告.md").write_text("# Hi\n\n![x](./assets/chart.png)\n", encoding="utf-8")
    (root / "assets" / "chart.png").write_bytes(_tiny_png())

    backend = ServerWorkspace(root=root, sandbox=SubprocessSandbox())
    out = await export_markdown_path(backend, "报告.md")
    assert out.output_path == "报告.docx"
    assert (root / "报告.docx").is_file()
    assert out.size_bytes == (root / "报告.docx").stat().st_size
    assert out.warnings == []


@pytest.mark.asyncio
async def test_export_markdown_path_missing_source(tmp_path: Path):
    backend = ServerWorkspace(root=tmp_path / "ws", sandbox=SubprocessSandbox())
    (tmp_path / "ws").mkdir()
    with pytest.raises(ExportMarkdownError, match="不存在"):
        await export_markdown_path(backend, "nope.md")
