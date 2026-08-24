"""Tests for attachment分流预解析 (``attachment_parse`` + persist hook + prompt)."""

from __future__ import annotations

import asyncio
import html
import sys
import time
import zipfile
from io import BytesIO
from pathlib import Path
from unittest.mock import patch

from agentcore.runtime.pipeline import _build_attachment_context
from agentcore.tools.sandbox.subprocess import SubprocessSandbox
from agentcore.workspace.attachment_parse import (
    ATTACHMENT_INLINE_MAX_CHARS,
    ParseStatus,
    extract_office_bytes,
    extract_office_payload,
    extract_table_preview,
    looks_like_scanned,
    parsed_copy_path,
    preparse_resident,
    should_preparse,
    should_preview_table,
    truncate_for_prompt,
)
from agentcore.workspace.attachments import persist_attachments
from agentcore.workspace.server import ServerWorkspace


def _ws(root: Path) -> ServerWorkspace:
    return ServerWorkspace(root=root, sandbox=SubprocessSandbox())


def _docx_bytes(*, paragraphs: list[str], table: list[list[str]] | None = None) -> bytes:
    from docx import Document

    doc = Document()
    for paragraph in paragraphs:
        doc.add_paragraph(paragraph)
    if table:
        grid = doc.add_table(rows=len(table), cols=max(len(row) for row in table))
        for i, row in enumerate(table):
            for j, value in enumerate(row):
                grid.cell(i, j).text = value
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _inprocess_extract():
    """Run the worker payload in-process so markitdown mocks still apply (pdf)."""

    def _run(data: bytes, ext: str, timeout: float, *, holder=None):
        return extract_office_payload(data, ext)

    return patch(
        "agentcore.workspace.attachment_parse._run_extract_subprocess",
        side_effect=_run,
    )


def test_should_preparse_routing():
    assert should_preparse("a.docx") is True
    assert should_preparse("a.pdf") is True
    assert should_preparse("a.pptx") is True
    assert should_preparse("notes.txt") is True
    assert should_preparse("a.xlsx") is False
    assert should_preparse("a.csv") is False
    assert should_preparse("a.xls") is False
    assert should_preparse("photo.png") is False
    assert should_preview_table("a.xlsx") is True
    assert should_preview_table("a.csv") is True
    assert should_preview_table("a.tsv") is True
    assert should_preview_table("a.xls") is True
    assert should_preview_table("a.pdf") is False


def test_looks_like_scanned_and_truncate():
    assert looks_like_scanned("", 100) is True
    assert looks_like_scanned("hi", 100) is True
    assert looks_like_scanned("a" * 50, 100) is False
    assert looks_like_scanned("a" * 100, 60_000) is True  # large + sparse
    assert looks_like_scanned("a" * 250, 60_000) is False

    body, clipped = truncate_for_prompt("x" * 100, limit=50)
    assert clipped is True
    assert len(body) == 50
    body2, clipped2 = truncate_for_prompt("short", limit=50)
    assert clipped2 is False
    assert body2 == "short"


def test_parsed_copy_path():
    assert parsed_copy_path("attachments/report.docx") == "attachments/report.docx.md"


async def test_extract_office_bytes_ok_and_skip(tmp_path: Path):
    payload = _docx_bytes(
        paragraphs=["Hello paragraph from python-docx with enough alphanumeric body."],
        table=[["Name", "Qty"], ["Widget", "42"]],
    )
    with patch(
        "agentcore.workspace.attachment_parse._convert_with_markitdown",
        side_effect=AssertionError("docx must not call markitdown"),
    ):
        ok = await extract_office_bytes(payload, ext=".docx")
    assert ok.status == ParseStatus.OK
    assert "Hello paragraph from python-docx" in ok.text
    assert "Widget" in ok.text
    assert "42" in ok.text
    assert "python-docx" in ok.detail

    skipped = await extract_office_bytes(b"PK", ext=".xlsx")
    assert skipped.status == ParseStatus.SKIPPED
    assert skipped.detail.startswith("skip_ext")


async def test_extract_docx_payload_skips_markitdown():
    payload = _docx_bytes(
        paragraphs=["In-process paragraph body with alphanumeric content here."]
    )
    with patch(
        "agentcore.workspace.attachment_parse._convert_with_markitdown",
        side_effect=AssertionError("docx payload must not call markitdown"),
    ):
        result = extract_office_payload(payload, ".docx")
    assert result.status == ParseStatus.OK
    assert "In-process paragraph body" in result.text


async def test_extract_pdf_payload_still_uses_markitdown():
    with patch(
        "agentcore.workspace.attachment_parse._convert_with_markitdown",
        return_value="PDF body with enough alphanumeric content for the scan heuristic.",
    ) as mocked:
        result = extract_office_payload(b"%PDF-fake", ".pdf")
    mocked.assert_called_once()
    assert result.status == ParseStatus.OK
    assert "PDF body" in result.text
    assert "markitdown" in result.detail


async def test_extract_hanging_worker_fails_within_budget_and_kills():
    from agentcore.workspace import attachment_parse as ap

    orig = ap._run_extract_subprocess
    seen: dict = {}

    def wrapped(data: bytes, ext: str, timeout: float, *, holder=None):
        result = orig(data, ext, timeout, holder=holder)
        seen["proc"] = None if holder is None else holder.get("proc")
        return result

    with (
        patch.object(ap, "_run_extract_subprocess", wrapped),
        patch.object(
            ap,
            "_extract_worker_argv",
            lambda ext: [
                sys.executable,
                "-c",
                "import sys,time; sys.stdin.buffer.read(); time.sleep(60)",
            ],
        ),
    ):
        t0 = time.monotonic()
        result = await extract_office_bytes(b"%PDF-x", ext=".pdf", timeout=0.4)
        elapsed = time.monotonic() - t0
    assert result.status == ParseStatus.FAILED
    assert "timeout" in result.detail
    assert elapsed < 2.0
    assert elapsed >= 0.3
    proc = seen.get("proc")
    assert proc is not None
    assert proc.poll() is not None


async def test_extract_does_not_block_event_loop():
    with patch(
        "agentcore.workspace.attachment_parse._extract_worker_argv",
        lambda ext: [
            sys.executable,
            "-c",
            "import sys,time; sys.stdin.buffer.read(); time.sleep(60)",
        ],
    ):
        ping: list[bool] = []

        async def _mark() -> None:
            await asyncio.sleep(0.05)
            ping.append(True)

        extract = asyncio.create_task(
            extract_office_bytes(b"%PDF-x", ext=".pdf", timeout=0.8)
        )
        marker = asyncio.create_task(_mark())
        done, _pending = await asyncio.wait({marker}, timeout=0.3)
        assert marker in done
        assert ping == [True]
        result = await extract
    assert result.status == ParseStatus.FAILED
    assert "timeout" in result.detail


async def test_preparse_docx_writes_md_copy(tmp_path: Path):
    ws = _ws(tmp_path)
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "brief.docx").write_bytes(
        _docx_bytes(
            paragraphs=["Hello from docx with enough alphanumeric body for the scan heuristic."]
        )
    )

    with patch(
        "agentcore.workspace.attachment_parse._convert_with_markitdown",
        side_effect=AssertionError("docx must not call markitdown"),
    ):
        result = await preparse_resident(
            ws, workspace_path="attachments/brief.docx", name="brief.docx"
        )

    assert result.status == ParseStatus.OK
    assert result.parsed_workspace_path == "attachments/brief.docx.md"
    assert "Hello from docx" in result.text
    assert "Hello from docx" in (
        tmp_path / "attachments" / "brief.docx.md"
    ).read_text(encoding="utf-8")


async def test_preparse_pdf_success_via_persist(tmp_path: Path):
    ws = _ws(tmp_path)
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "paper.pdf").write_bytes(b"%PDF-fake")

    with _inprocess_extract(), patch(
        "agentcore.workspace.attachment_parse._convert_with_markitdown",
        return_value="Abstract\n\nThis paper studies agents." + ("x" * 20),
    ):
        out = await persist_attachments(
            ws,
            [
                {
                    "name": "paper.pdf",
                    "path": "attachments/paper.pdf",
                    "text": "",
                    "binary": True,
                    "workspace_path": "attachments/paper.pdf",
                }
            ],
        )

    assert out[0]["parse_status"] == "ok"
    assert out[0]["parsed_workspace_path"] == "attachments/paper.pdf.md"
    assert "This paper studies agents" in out[0]["text"]
    assert (tmp_path / "attachments" / "paper.pdf.md").exists()


async def test_preparse_scanned_pdf_writes_notice(tmp_path: Path):
    ws = _ws(tmp_path)
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "scan.pdf").write_bytes(b"%PDF" + b"\x00" * 100)

    with _inprocess_extract(), patch(
        "agentcore.workspace.attachment_parse._convert_with_markitdown",
        return_value="   \n",  # empty text layer
    ):
        result = await preparse_resident(
            ws, workspace_path="attachments/scan.pdf", name="scan.pdf"
        )

    assert result.status == ParseStatus.SCANNED
    assert "scanned" in result.text.lower() or "OCR" in result.text
    assert result.parsed_workspace_path == "attachments/scan.pdf.md"
    note = (tmp_path / "attachments" / "scan.pdf.md").read_text(encoding="utf-8")
    assert "OCR" in note

async def test_preparse_xlsx_invalid_has_no_body(tmp_path: Path):
    ws = _ws(tmp_path)
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "report.xlsx").write_bytes(b"PK\x03\x04")

    out = await persist_attachments(
        ws,
        [
            {
                "name": "report.xlsx",
                "path": "attachments/report.xlsx",
                "text": "",
                "binary": True,
                "workspace_path": "attachments/report.xlsx",
            }
        ],
    )
    assert out[0]["parse_status"] == "failed"
    assert "parsed_workspace_path" not in out[0]
    assert "table_preview" not in out[0]
    assert not (out[0].get("text") or "").strip()
    assert not (tmp_path / "attachments" / "report.xlsx.md").exists()
    # Original untouched.
    assert (tmp_path / "attachments" / "report.xlsx").read_bytes() == b"PK\x03\x04"


async def test_preparse_failure_falls_back(tmp_path: Path):
    ws = _ws(tmp_path)
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "broken.docx").write_bytes(b"not-a-docx")

    out = await persist_attachments(
        ws,
        [
            {
                "name": "broken.docx",
                "path": "attachments/broken.docx",
                "text": "",
                "binary": True,
                "workspace_path": "attachments/broken.docx",
            }
        ],
    )

    assert out[0]["parse_status"] == "failed"
    assert not (out[0].get("text") or "").strip()
    assert out[0]["workspace_path"] == "attachments/broken.docx"
    assert not (tmp_path / "attachments" / "broken.docx.md").exists()

    # Prompt steers file_read for office/PDF — not code_execute.
    ctx = await _build_attachment_context(out)
    assert ctx is not None
    assert "[binary / office-pdf]" in ctx
    assert "file_read" in ctx
    assert "do not default to code_execute" in ctx
    assert "openpyxl" not in ctx

async def test_context_preparsed_inline_and_large_truncation():
    small = {
        "name": "a.docx",
        "path": "attachments/a.docx",
        "binary": True,
        "workspace_path": "attachments/a.docx",
        "parsed_workspace_path": "attachments/a.docx.md",
        "parse_status": "ok",
        "text": "Hello world from docx extract.",
    }
    out = await _build_attachment_context([small])
    assert out is not None
    assert "Hello world from docx extract" in out
    assert "pre-parsed → attachments/a.docx.md" in out
    assert "pre-parsed at upload" in out
    assert "lossy for tabular content" in out
    assert "Do not use this extract as the data source" in out

    huge_text = "Z" * (ATTACHMENT_INLINE_MAX_CHARS + 500)
    large = {
        **small,
        "name": "big.docx",
        "text": huge_text,
        "parsed_workspace_path": "attachments/big.docx.md",
        "workspace_path": "attachments/big.docx",
    }
    out2 = await _build_attachment_context([large])
    assert out2 is not None
    assert "truncated" in out2
    assert "full extracted text is at attachments/big.docx.md" in out2
    # Inline body capped.
    assert "Z" * (ATTACHMENT_INLINE_MAX_CHARS + 1) not in out2


async def test_context_scanned_shows_notice():
    out = await _build_attachment_context(
        [
            {
                "name": "scan.pdf",
                "path": "attachments/scan.pdf",
                "binary": True,
                "workspace_path": "attachments/scan.pdf",
                "parsed_workspace_path": "attachments/scan.pdf.md",
                "parse_status": "scanned",
                "text": "This file appears to be a scanned / image-only document.",
            }
        ]
    )
    assert out is not None
    assert "scanned / no text layer" in out
    assert "image-only" in out


async def test_plain_txt_binary_resident_decodes(tmp_path: Path):
    ws = _ws(tmp_path)
    (tmp_path / "attachments").mkdir()
    (tmp_path / "attachments" / "notes.txt").write_text("plain notes\n", encoding="utf-8")

    result = await preparse_resident(
        ws, workspace_path="attachments/notes.txt", name="notes.txt"
    )
    assert result.status == ParseStatus.OK
    assert result.parsed_workspace_path == "attachments/notes.txt"
    assert "plain notes" in result.text
    # No redundant .md for already-text originals.
    assert not (tmp_path / "attachments" / "notes.txt.md").exists()


async def test_preparse_read_failure_is_failed_not_raise(tmp_path: Path):
    ws = _ws(tmp_path)
    # File missing → read_bytes raises WorkspaceError subclass.
    result = await preparse_resident(
        ws, workspace_path="attachments/missing.docx", name="missing.docx"
    )
    assert result.status == ParseStatus.FAILED
    assert "read" in result.detail


_SECRET_TAIL = "UNIQUE_TAIL_ROW_SHOULD_NOT_ENTER_PROMPT"


def _wide_csv(*, rows: int = 20) -> str:
    lines = ["date,amount,memo"]
    for i in range(1, rows + 1):
        memo = _SECRET_TAIL if i == rows else f"item-{i}"
        lines.append(f"2024-01-{i:02d},{i * 1.5:.1f},{memo}")
    return "\n".join(lines) + "\n"


def _col_letter(index: int) -> str:
    text = ""
    number = index + 1
    while number:
        number, rem = divmod(number - 1, 26)
        text = chr(65 + rem) + text
    return text


def _minimal_xlsx(headers: list[str], rows: list[list[object]]) -> bytes:
    ns = "http://schemas.openxmlformats.org/spreadsheetml/2006/main"
    rel = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
    pkg = "http://schemas.openxmlformats.org/package/2006/relationships"
    ct = "http://schemas.openxmlformats.org/package/2006/content-types"

    def cell(ref: str, value: object) -> str:
        if isinstance(value, int | float) and not isinstance(value, bool):
            return f'<c r="{ref}"><v>{value}</v></c>'
        escaped = html.escape(str(value), quote=True)
        return f'<c r="{ref}" t="inlineStr"><is><t>{escaped}</t></is></c>'

    sheet_rows: list[str] = []
    for r_i, row in enumerate([headers, *rows], start=1):
        cells = "".join(cell(f"{_col_letter(c)}{r_i}", value) for c, value in enumerate(row))
        sheet_rows.append(f'<row r="{r_i}">{cells}</row>')
    sheet = (
        f'<?xml version="1.0" encoding="UTF-8"?>'
        f'<worksheet xmlns="{ns}"><sheetData>{"".join(sheet_rows)}</sheetData></worksheet>'
    )
    workbook = (
        f'<?xml version="1.0" encoding="UTF-8"?>'
        f'<workbook xmlns="{ns}" xmlns:r="{rel}">'
        f'<sheets><sheet name="Sheet1" sheetId="1" r:id="rId1"/></sheets>'
        f"</workbook>"
    )
    rels = (
        f'<?xml version="1.0" encoding="UTF-8"?>'
        f'<Relationships xmlns="{pkg}">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" '
        'Target="xl/workbook.xml"/>'
        "</Relationships>"
    )
    wb_rels = (
        f'<?xml version="1.0" encoding="UTF-8"?>'
        f'<Relationships xmlns="{pkg}">'
        '<Relationship Id="rId1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/worksheet" '
        'Target="worksheets/sheet1.xml"/>'
        "</Relationships>"
    )
    ctypes = (
        f'<?xml version="1.0" encoding="UTF-8"?>'
        f'<Types xmlns="{ct}">'
        '<Default Extension="rels" '
        'ContentType="application/vnd.openxmlformats-package.relationships+xml"/>'
        '<Default Extension="xml" ContentType="application/xml"/>'
        '<Override PartName="/xl/workbook.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/>'
        '<Override PartName="/xl/worksheets/sheet1.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.worksheet+xml"/>'
        "</Types>"
    )
    buf = BytesIO()
    with zipfile.ZipFile(buf, "w") as zf:
        zf.writestr("[Content_Types].xml", ctypes)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("xl/workbook.xml", workbook)
        zf.writestr("xl/_rels/workbook.xml.rels", wb_rels)
        zf.writestr("xl/worksheets/sheet1.xml", sheet)
    return buf.getvalue()


def test_extract_table_preview_csv_structure_hides_tail():
    body = _wide_csv(rows=20)
    result = extract_table_preview(body.encode("utf-8"), ".csv")
    assert result.status == ParseStatus.OK
    assert result.preview is not None
    sheet = result.preview.sheets[0]
    assert [col.name for col in sheet.columns] == ["date", "amount", "memo"]
    assert sheet.row_count == 20
    types = {col.name: col.inferred_type for col in sheet.columns}
    assert types["date"] == "date"
    assert types["amount"] == "float"
    assert types["memo"] == "string"
    assert len(sheet.sample_rows) == 5
    dumped = str(result.preview.to_dict())
    assert _SECRET_TAIL not in dumped
    assert "item-1" in dumped


def test_extract_table_preview_xlsx_structure_hides_tail():
    rows: list[list[object]] = []
    for i in range(1, 21):
        memo = _SECRET_TAIL if i == 20 else f"item-{i}"
        rows.append([f"2024-01-{i:02d}", i * 1.5, memo])
    result = extract_table_preview(_minimal_xlsx(["date", "amount", "memo"], rows), ".xlsx")
    assert result.status == ParseStatus.OK
    assert result.preview is not None
    sheet = result.preview.sheets[0]
    assert sheet.row_count == 20
    assert [col.name for col in sheet.columns] == ["date", "amount", "memo"]
    dumped = str(result.preview.to_dict())
    assert _SECRET_TAIL not in dumped


def test_extract_table_preview_xls_unsupported():
    result = extract_table_preview(b"\xd0\xcf\x11\xe0" + b"\x00" * 20, ".xls")
    assert result.status == ParseStatus.FAILED
    assert result.preview is None
    assert "xls_unsupported" in result.detail


async def test_persist_csv_structure_does_not_inline_full_table(tmp_path: Path):
    ws = _ws(tmp_path)
    body = _wide_csv(rows=20)
    out = await persist_attachments(
        ws, [{"name": "ledger.csv", "path": "/local/ledger.csv", "text": body}]
    )
    assert out[0]["workspace_path"] == "attachments/ledger.csv"
    assert out[0]["parse_status"] == "ok"
    assert not (out[0].get("text") or "").strip()
    preview = out[0]["table_preview"]
    assert preview["sheets"][0]["row_count"] == 20
    assert [c["name"] for c in preview["sheets"][0]["columns"]] == ["date", "amount", "memo"]
    assert _SECRET_TAIL not in str(preview)
    on_disk = (tmp_path / "attachments" / "ledger.csv").read_text(encoding="utf-8")
    assert _SECRET_TAIL in on_disk

    ctx = await _build_attachment_context(out, available_tools=frozenset())
    assert ctx is not None
    assert "[table / structure]" in ctx
    assert "rows: 20" in ctx
    assert "date:date" in ctx
    assert "amount:float" in ctx
    assert _SECRET_TAIL not in ctx
    assert "includes code_execute" not in ctx
    assert "with code_execute" not in ctx


async def test_persist_xlsx_structure_does_not_inline_full_table(tmp_path: Path):
    ws = _ws(tmp_path)
    (tmp_path / "attachments").mkdir()
    rows: list[list[object]] = []
    for i in range(1, 21):
        memo = _SECRET_TAIL if i == 20 else f"item-{i}"
        rows.append([f"2024-01-{i:02d}", i * 1.5, memo])
    raw = _minimal_xlsx(["date", "amount", "memo"], rows)
    (tmp_path / "attachments" / "ledger.xlsx").write_bytes(raw)
    out = await persist_attachments(
        ws,
        [
            {
                "name": "ledger.xlsx",
                "path": "attachments/ledger.xlsx",
                "text": "",
                "binary": True,
                "workspace_path": "attachments/ledger.xlsx",
            }
        ],
    )
    assert out[0]["parse_status"] == "ok"
    assert not (out[0].get("text") or "").strip()
    assert _SECRET_TAIL not in str(out[0]["table_preview"])
    ctx = await _build_attachment_context(out, available_tools=frozenset())
    assert ctx is not None
    assert "[table / structure]" in ctx
    assert "rows: 20" in ctx
    assert _SECRET_TAIL not in ctx
    assert "includes code_execute" not in ctx
    assert "with code_execute" not in ctx
